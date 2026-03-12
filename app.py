# app.py
# ===============================
# SISTEMA SATELITAL DE ANÁLISIS AMBIENTAL INTEGRAL
# Carbono + Biodiversidad + Análisis Forrajero
# Con mapas continuos, dashboard interactivo e informe con IA (Gemini)
# ===============================

# ✅ ABSOLUTAMENTE PRIMERO: Importar streamlit
import streamlit as st
# ✅ LUEGO: Configurar la página
st.set_page_config(
    page_title="Sistema Satelital de Análisis Ambiental Integral",
    page_icon="🌎",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ===== IMPORTS ESTÁNDAR =====
import pandas as pd
import numpy as np
import tempfile
import os
import zipfile
import math
from math import log
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
import plotly.figure_factory as ff
from plotly.subplots import make_subplots
from io import BytesIO, StringIO
from datetime import datetime, timedelta
import json
import base64
import warnings
import requests
import xml.etree.ElementTree as ET
from typing import Optional, Dict, Any, List, Tuple
import random

# ===== IMPORTACIÓN DE MÓDULOS IA =====
from modules.ia_integration import (
    preparar_resumen,
    generar_analisis_carbono,
    generar_analisis_biodiversidad,
    generar_analisis_espectral,
    generar_analisis_forrajero,          # nuevo
    generar_recomendaciones_integradas
)

# ===== IMPORTACIONES GOOGLE EARTH ENGINE =====
try:
    import ee
    GEE_AVAILABLE = True
except ImportError:
    GEE_AVAILABLE = False
    st.warning("⚠️ Google Earth Engine no está instalado. Para usar datos satelitales reales, instala con: pip install earthengine-api")

warnings.filterwarnings('ignore')

# ===== LIBRERÍAS GEOESPACIALES =====
import folium
from streamlit_folium import st_folium, folium_static
from folium.plugins import Fullscreen, MousePosition, HeatMap
import geopandas as gpd
from shapely.geometry import Polygon, Point, shape, MultiPolygon
from shapely.ops import unary_union
import pyproj
from branca.colormap import LinearColormap
import matplotlib.cm as cm
from scipy.interpolate import griddata
from matplotlib.colors import LinearSegmentedColormap

# ===== CONFIGURACIÓN DE IA (GEMINI) =====
GEMINI_API_KEY = st.secrets.get("GEMINI_API_KEY", os.getenv("GEMINI_API_KEY"))
if not GEMINI_API_KEY:
    st.warning("⚠️ No se encontró API Key de Gemini. La IA no estará disponible.")
else:
    os.environ["GEMINI_API_KEY"] = GEMINI_API_KEY

# ===== INICIALIZACIÓN DE GOOGLE EARTH ENGINE =====
def inicializar_gee():
    if not GEE_AVAILABLE:
        return False
    try:
        gee_secret = os.environ.get('GEE_SERVICE_ACCOUNT')
        if gee_secret:
            try:
                credentials_info = json.loads(gee_secret.strip())
                credentials = ee.ServiceAccountCredentials(
                    credentials_info['client_email'],
                    key_data=json.dumps(credentials_info)
                )
                ee.Initialize(credentials, project='ee-mawucano25')
                st.session_state.gee_authenticated = True
                st.session_state.gee_project = 'ee-mawucano25'
                return True
            except Exception as e:
                print(f"⚠️ Error con Service Account: {str(e)}")
        try:
            ee.Initialize(project='ee-mawucano25')
            st.session_state.gee_authenticated = True
            st.session_state.gee_project = 'ee-mawucano25'
            return True
        except Exception as e:
            print(f"⚠️ Error inicialización local: {str(e)}")
        st.session_state.gee_authenticated = False
        return False
    except Exception as e:
        st.session_state.gee_authenticated = False
        print(f"❌ Error crítico GEE: {str(e)}")
        return False

# ===== LIBRERÍAS PARA REPORTES =====
try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, letter, landscape
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle,
        PageBreak, KeepTogether, PageTemplate, Frame, NextPageTemplate,
        BaseDocTemplate, FrameBreak
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    REPORTPDF_AVAILABLE = True
except ImportError:
    REPORTPDF_AVAILABLE = False
    st.warning("ReportLab no está instalado. La generación de PDFs estará limitada.")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    REPORTDOCX_AVAILABLE = True
except ImportError:
    REPORTDOCX_AVAILABLE = False
    st.warning("python-docx no está instalado. La generación de DOCX estará limitada.")

# ===============================
# 🌦️ CONECTOR CLIMÁTICO TROPICAL
# ===============================
class ConectorClimaticoTropical:
    def __init__(self):
        pass

    def obtener_datos_climaticos(self, lat: float, lon: float) -> Dict:
        if -5 <= lat <= 5 and -75 <= lon <= -50:  # Amazonía central
            return {'precipitacion': 2500 + random.uniform(-200, 200), 'temperatura': 26 + random.uniform(-1, 1)}
        elif abs(lat) < 10 and -82 <= lon <= -75:  # Chocó
            return {'precipitacion': 4000 + random.uniform(-300, 300), 'temperatura': 27 + random.uniform(-1, 1)}
        elif -15 <= lat < -5 and -70 <= lon <= -50:  # Sur amazónico
            return {'precipitacion': 1800 + random.uniform(-200, 200), 'temperatura': 25 + random.uniform(-1, 1)}
        elif -34 <= lat <= -22 and -73 <= lon <= -53:  # Argentina templada
            return {'precipitacion': 800 + random.uniform(-100, 100), 'temperatura': 18 + random.uniform(-2, 2)}
        else:
            return {'precipitacion': 1200 + random.uniform(-200, 200), 'temperatura': 22 + random.uniform(-2, 2)}

# ===============================
# 🌳 METODOLOGÍA VERRA (ajustada para cultivos)
# ===============================
class MetodologiaVerra:
    def __init__(self):
        self.factores = {
            'conversion_carbono': 0.47,
            'ratio_co2': 3.67,
            'ratio_raiz': 0.24,
            'proporcion_madera_muerta': 0.15,
            'acumulacion_hojarasca': 5.0,
            'carbono_suelo': 2.5
        }
        self.factores_vegetacion = {
            'amazonia': {'factor_biomasa': 1.2, 'factor_suelo': 1.0, 'factor_madera': 1.0},
            'choco': {'factor_biomasa': 1.3, 'factor_suelo': 1.1, 'factor_madera': 1.0},
            'seco': {'factor_biomasa': 0.8, 'factor_suelo': 0.7, 'factor_madera': 0.8},
            'vid': {'factor_biomasa': 0.15, 'factor_suelo': 0.6, 'factor_madera': 0.05},
            'cultivo': {'factor_biomasa': 0.2, 'factor_suelo': 0.7, 'factor_madera': 0.1},
            'agricola': {'factor_biomasa': 0.25, 'factor_suelo': 0.8, 'factor_madera': 0.1},
            'pampa': {'factor_biomasa': 0.4, 'factor_suelo': 0.9, 'factor_madera': 0.2},
            'andes': {'factor_biomasa': 0.6, 'factor_suelo': 0.9, 'factor_madera': 0.5}
        }

    def calcular_carbono_hectarea(self, ndvi: float, tipo_bosque: str, precipitacion: float) -> Dict:
        factores_veg = self.factores_vegetacion.get(tipo_bosque, 
            {'factor_biomasa': 1.0, 'factor_suelo': 1.0, 'factor_madera': 1.0})
        
        if tipo_bosque in ['vid', 'cultivo', 'agricola']:
            factor_precip = min(1.3, max(0.7, precipitacion / 1500))
        else:
            factor_precip = min(2.0, max(0.5, precipitacion / 1500))
        
        if tipo_bosque in ['vid', 'cultivo', 'agricola']:
            if ndvi > 0.7:
                agb_ton_ha = (30 + (ndvi - 0.7) * 50) * factor_precip
            elif ndvi > 0.5:
                agb_ton_ha = (20 + (ndvi - 0.5) * 60) * factor_precip
            elif ndvi > 0.3:
                agb_ton_ha = (10 + (ndvi - 0.3) * 50) * factor_precip
            else:
                agb_ton_ha = (5 + ndvi * 30) * factor_precip
        else:
            if ndvi > 0.7:
                agb_ton_ha = (150 + (ndvi - 0.7) * 300) * factor_precip
            elif ndvi > 0.5:
                agb_ton_ha = (80 + (ndvi - 0.5) * 350) * factor_precip
            elif ndvi > 0.3:
                agb_ton_ha = (30 + (ndvi - 0.3) * 250) * factor_precip
            else:
                agb_ton_ha = (5 + ndvi * 100) * factor_precip
        
        agb_ton_ha *= factores_veg['factor_biomasa']
        if tipo_bosque == "vid":
            agb_ton_ha *= 0.9
        elif tipo_bosque == "cultivo":
            agb_ton_ha *= 0.8
        
        carbono_agb = agb_ton_ha * self.factores['conversion_carbono']
        
        if tipo_bosque in ['vid', 'cultivo', 'agricola']:
            carbono_bgb = carbono_agb * (self.factores['ratio_raiz'] * 0.7)
        else:
            carbono_bgb = carbono_agb * self.factores['ratio_raiz']
        
        carbono_dw = carbono_agb * self.factores['proporcion_madera_muerta'] * factores_veg['factor_madera']
        
        if tipo_bosque in ['vid', 'cultivo', 'agricola']:
            carbono_li = self.factores['acumulacion_hojarasca'] * 0.3 * self.factores['conversion_carbono']
        else:
            carbono_li = self.factores['acumulacion_hojarasca'] * self.factores['conversion_carbono']
        
        carbono_soc = self.factores['carbono_suelo'] * factores_veg['factor_suelo']
        
        carbono_total = carbono_agb + carbono_bgb + carbono_dw + carbono_li + carbono_soc
        co2_equivalente = carbono_total * self.factores['ratio_co2']
        
        return {
            'carbono_total_ton_ha': round(carbono_total, 2),
            'co2_equivalente_ton_ha': round(co2_equivalente, 2),
            'biomasa_aerea_ton_ha': round(agb_ton_ha, 2),
            'desglose': {
                'AGB': round(carbono_agb, 2),
                'BGB': round(carbono_bgb, 2),
                'DW': round(carbono_dw, 2),
                'LI': round(carbono_li, 2),
                'SOC': round(carbono_soc, 2)
            },
            'tipo_vegetacion': tipo_bosque
        }

# ===============================
# 🦋 ANÁLISIS DE BIODIVERSIDAD
# ===============================
class AnalisisBiodiversidad:
    def __init__(self):
        self.parametros = {
            'amazonia': {'riqueza_base': 150, 'abundancia_base': 1000, 'factor_ndvi': 0.8, 'es_cultivo': False},
            'choco': {'riqueza_base': 120, 'abundancia_base': 800, 'factor_ndvi': 0.8, 'es_cultivo': False},
            'andes': {'riqueza_base': 100, 'abundancia_base': 600, 'factor_ndvi': 0.8, 'es_cultivo': False},
            'pampa': {'riqueza_base': 50, 'abundancia_base': 300, 'factor_ndvi': 0.8, 'es_cultivo': False},
            'seco': {'riqueza_base': 40, 'abundancia_base': 200, 'factor_ndvi': 0.8, 'es_cultivo': False},
            'cultivo': {'riqueza_base': 10, 'abundancia_base': 50, 'factor_ndvi': 0.2, 'es_cultivo': True},
            'vid': {'riqueza_base': 8, 'abundancia_base': 40, 'factor_ndvi': 0.1, 'es_cultivo': True},
            'agricola': {'riqueza_base': 15, 'abundancia_base': 60, 'factor_ndvi': 0.3, 'es_cultivo': True}
        }

    def calcular_shannon(self, ndvi: float, tipo_ecosistema: str, area_ha: float, precipitacion: float) -> Dict:
        params = self.parametros.get(tipo_ecosistema, {'riqueza_base': 60, 'abundancia_base': 400, 'factor_ndvi': 0.5, 'es_cultivo': False})
        factor_ndvi = 1.0 + (ndvi * params['factor_ndvi'])
        if params['es_cultivo']:
            factor_area = min(1.3, math.log10(area_ha + 1) * 0.2 + 1)
        else:
            factor_area = min(2.0, math.log10(area_ha + 1) * 0.5 + 1)
        if tipo_ecosistema in ['amazonia', 'choco']:
            factor_precip = min(1.5, precipitacion / 2000)
        elif params['es_cultivo']:
            factor_precip = 1.0 + (precipitacion / 2000 * 0.3)
        else:
            factor_precip = 1.0
        
        riqueza_especies = int(params['riqueza_base'] * factor_ndvi * factor_area * factor_precip * random.uniform(0.8, 1.2))
        if params['es_cultivo']:
            abundancia_total = int(params['abundancia_base'] * factor_ndvi * factor_area * factor_precip * random.uniform(0.9, 1.1))
        else:
            abundancia_total = int(params['abundancia_base'] * factor_ndvi * factor_area * factor_precip * random.uniform(0.9, 1.1))
        
        especies = []
        abundancia_acumulada = 0
        if params['es_cultivo']:
            if riqueza_especies > 0:
                abundancia_principal = int(abundancia_total * random.uniform(0.7, 0.9))
                especies.append({'especie_id': 1, 'abundancia': abundancia_principal, 'nombre': tipo_ecosistema.capitalize()})
                abundancia_acumulada += abundancia_principal
                for i in range(2, riqueza_especies + 1):
                    abundancia = int((abundancia_total - abundancia_principal) / max(riqueza_especies - 1, 1) * random.uniform(0.5, 1.5))
                    if abundancia > 0:
                        especies.append({'especie_id': i, 'abundancia': abundancia, 'nombre': f'Especie {i}'})
                        abundancia_acumulada += abundancia
        else:
            for i in range(1, riqueza_especies + 1):
                abundancia = int((abundancia_total / max(riqueza_especies, 1)) * random.lognormvariate(0, 0.5))
                if abundancia > 0:
                    especies.append({'especie_id': i, 'abundancia': abundancia, 'nombre': f'Especie {i}'})
                    abundancia_acumulada += abundancia
        
        for especie in especies:
            especie['proporcion'] = especie['abundancia'] / abundancia_acumulada if abundancia_acumulada > 0 else 0
        
        shannon = 0
        for especie in especies:
            if especie['proporcion'] > 0:
                shannon -= especie['proporcion'] * math.log(especie['proporcion'])
        
        if params['es_cultivo']:
            if shannon > 1.5:
                categoria = "Alta (para cultivo)"
                color = "#3b82f6"
            elif shannon > 1.0:
                categoria = "Moderada (para cultivo)"
                color = "#f59e0b"
            elif shannon > 0.5:
                categoria = "Baja (típico de monocultivo)"
                color = "#ef4444"
            else:
                categoria = "Muy Baja (monocultivo puro)"
                color = "#991b1b"
        else:
            if shannon > 3.5:
                categoria = "Muy Alta"
                color = "#10b981"
            elif shannon > 2.5:
                categoria = "Alta"
                color = "#3b82f6"
            elif shannon > 1.5:
                categoria = "Moderada"
                color = "#f59e0b"
            elif shannon > 0.5:
                categoria = "Baja"
                color = "#ef4444"
            else:
                categoria = "Muy Baja"
                color = "#991b1b"
        
        return {
            'indice_shannon': round(shannon, 3),
            'categoria': categoria,
            'color': color,
            'riqueza_especies': riqueza_especies,
            'abundancia_total': abundancia_acumulada,
            'especies_muestra': especies[:10],
            'es_cultivo': params['es_cultivo']
        }

# ===============================
# 🐮 ANÁLISIS FORRAJERO (completo)
# ===============================
class AnalisisForrajero:
    def __init__(self):
        self.parametros_forrajeros = {
            'pastizal_natural': {
                'productividad_kg_ms_ha': {'bajo': 2000, 'medio': 4000, 'alto': 6000},
                'eficiencia_aprovechamiento': 0.5,
                'tasa_crecimiento_diario': {'bajo': 15, 'medio': 30, 'alto': 45},
                'densidad_forraje': 2.5
            },
            'pastura_mejorada': {
                'productividad_kg_ms_ha': {'bajo': 4000, 'medio': 8000, 'alto': 12000},
                'eficiencia_aprovechamiento': 0.6,
                'tasa_crecimiento_diario': {'bajo': 25, 'medio': 50, 'alto': 75},
                'densidad_forraje': 3.0
            },
            'silvopastoril': {
                'productividad_kg_ms_ha': {'bajo': 3000, 'medio': 6000, 'alto': 9000},
                'eficiencia_aprovechamiento': 0.55,
                'tasa_crecimiento_diario': {'bajo': 20, 'medio': 40, 'alto': 60},
                'densidad_forraje': 2.8
            },
            'agroforestal': {
                'productividad_kg_ms_ha': {'bajo': 2500, 'medio': 5000, 'alto': 7500},
                'eficiencia_aprovechamiento': 0.45,
                'tasa_crecimiento_diario': {'bajo': 18, 'medio': 36, 'alto': 54},
                'densidad_forraje': 2.6
            }
        }
        self.consumo_animal = {
            'vaca_adulta': 12,
            'novillo': 10,
            'ternero': 4,
            'vaca_secas': 8,
            'vaca_lactancia': 14,
            'equivalente_vaca': 12
        }
        self.factores_ndvi = {
            'bajo': {'ndvi_min': -1.0, 'ndvi_max': 0.2, 'factor': 0.3},
            'medio': {'ndvi_min': 0.2, 'ndvi_max': 0.5, 'factor': 0.6},
            'alto': {'ndvi_min': 0.5, 'ndvi_max': 1.0, 'factor': 1.0}
        }

    def estimar_disponibilidad_forrajera(self, ndvi: float, tipo_sistema: str, area_ha: float) -> Dict:
        if ndvi < 0.2:
            categoria_productividad = 'bajo'
        elif ndvi > 0.5:
            categoria_productividad = 'alto'
        else:
            categoria_productividad = 'medio'
        
        params = self.parametros_forrajeros.get(tipo_sistema, self.parametros_forrajeros['pastizal_natural'])
        productividad_base = params['productividad_kg_ms_ha'][categoria_productividad]
        factor_ndvi = 0.5 + (ndvi * 0.5)
        productividad_ajustada = productividad_base * factor_ndvi * random.uniform(0.9, 1.1)
        disponibilidad_total_kg_ms = productividad_ajustada * area_ha
        forraje_aprovechable_kg_ms = disponibilidad_total_kg_ms * params['eficiencia_aprovechamiento']
        tasa_crecimiento = params['tasa_crecimiento_diario'][categoria_productividad] * area_ha
        
        return {
            'productividad_kg_ms_ha': round(productividad_ajustada, 2),
            'disponibilidad_total_kg_ms': round(disponibilidad_total_kg_ms, 2),
            'forraje_aprovechable_kg_ms': round(forraje_aprovechable_kg_ms, 2),
            'tasa_crecimiento_diario_kg': round(tasa_crecimiento, 2),
            'categoria_productividad': categoria_productividad,
            'densidad_forraje_kg_m3': params['densidad_forraje']
        }

    def calcular_equivalentes_vaca(self, forraje_aprovechable_kg_ms: float, dias_permanencia: int = 1) -> Dict:
        consumo_ev_diario = self.consumo_animal['equivalente_vaca']
        ev_por_dia = forraje_aprovechable_kg_ms / consumo_ev_diario
        ev_para_periodo = forraje_aprovechable_kg_ms / (consumo_ev_diario * dias_permanencia)
        consumo_total_periodo = ev_para_periodo * consumo_ev_diario * dias_permanencia
        margen_seguridad = 0.8
        return {
            'ev_por_dia': round(ev_por_dia, 2),
            'ev_para_periodo': round(ev_para_periodo, 2),
            'ev_recomendado': round(ev_para_periodo * margen_seguridad, 2),
            'consumo_ev_diario_kg': consumo_ev_diario,
            'consumo_total_periodo_kg': round(consumo_total_periodo, 2),
            'dias_permanencia': dias_permanencia,
            'margen_seguridad': '20%'
        }

    def calcular_dias_permanencia(self, forraje_aprovechable_kg_ms: float, num_ev: float) -> Dict:
        consumo_ev_diario = self.consumo_animal['equivalente_vaca']
        consumo_diario_total = num_ev * consumo_ev_diario
        dias_permanencia_basico = forraje_aprovechable_kg_ms / consumo_diario_total
        dias_permanencia_ajustado = dias_permanencia_basico * 1.2
        dias_recomendados = min(30, int(dias_permanencia_ajustado))
        return {
            'dias_basico': round(dias_permanencia_basico, 1),
            'dias_ajustado': round(dias_permanencia_ajustado, 1),
            'dias_recomendados': dias_recomendados,
            'consumo_diario_total_kg': round(consumo_diario_total, 2),
            'forraje_disponible_kg': round(forraje_aprovechable_kg_ms, 2),
            'num_ev': num_ev
        }

    def dividir_lote_en_sublotes(self, area_total_ha: float, disponibilidad_forrajera_kg_ms_ha: float, heterogeneidad: float = 0.3) -> List[Dict]:
        if area_total_ha < 10:
            num_sublotes = 2
        elif area_total_ha < 50:
            num_sublotes = 3
        elif area_total_ha < 100:
            num_sublotes = 4
        else:
            num_sublotes = min(6, int(area_total_ha / 20))
        sublotes = []
        area_por_sublote = area_total_ha / num_sublotes
        for i in range(num_sublotes):
            variacion = 1 + random.uniform(-heterogeneidad, heterogeneidad)
            disponibilidad_sublote = disponibilidad_forrajera_kg_ms_ha * variacion
            forraje_sublote_kg_ms = disponibilidad_sublote * area_por_sublote
            forraje_aprovechable = forraje_sublote_kg_ms * 0.5
            sublotes.append({
                'sublote_id': i + 1,
                'area_ha': round(area_por_sublote, 2),
                'disponibilidad_kg_ms_ha': round(disponibilidad_sublote, 2),
                'forraje_total_kg_ms': round(forraje_sublote_kg_ms, 2),
                'forraje_aprovechable_kg_ms': round(forraje_aprovechable, 2),
                'productividad_relativa': round(variacion, 2)
            })
        return sublotes

    def generar_recomendaciones_rotacion(self, sublotes: List[Dict], num_ev_total: float) -> Dict:
        forraje_total_aprovechable = sum(s['forraje_aprovechable_kg_ms'] for s in sublotes)
        consumo_diario_total = num_ev_total * self.consumo_animal['equivalente_vaca']
        dias_rotacion_total = forraje_total_aprovechable / consumo_diario_total
        plan_rotacion = []
        for sublote in sublotes:
            dias_en_sublote = int((sublote['forraje_aprovechable_kg_ms'] / consumo_diario_total) * 0.8)
            dias_descanso = dias_en_sublote * 3
            plan_rotacion.append({
                'sublote': sublote['sublote_id'],
                'area_ha': sublote['area_ha'],
                'dias_uso': max(3, dias_en_sublote),
                'dias_descanso': max(21, dias_descanso),
                'productividad': sublote['productividad_relativa'],
                'recomendacion': self._generar_recomendacion_sublote(sublote['productividad_relativa'])
            })
        dias_ciclo = sum(p['dias_uso'] + p['dias_descanso'] for p in plan_rotacion) / len(plan_rotacion)
        return {
            'forraje_total_aprovechable_kg': round(forraje_total_aprovechable, 2),
            'consumo_diario_total_kg': round(consumo_diario_total, 2),
            'dias_rotacion_total': round(dias_rotacion_total, 1),
            'num_ev': num_ev_total,
            'plan_rotacion': plan_rotacion,
            'dias_ciclo_promedio': round(dias_ciclo, 1),
            'intensidad_rotacion': self._clasificar_intensidad_rotacion(dias_ciclo)
        }

    def _generar_recomendacion_sublote(self, productividad: float) -> str:
        if productividad > 1.2:
            return "Alta productividad - Considerar manejo intensivo con pastoreo rotativo"
        elif productividad > 0.8:
            return "Productividad media - Ideal para rotación estándar"
        else:
            return "Baja productividad - Requiere recuperación, considerar descanso prolongado"

    def _clasificar_intensidad_rotacion(self, dias_ciclo: float) -> str:
        if dias_ciclo < 30:
            return "Alta intensidad - Rotación rápida"
        elif dias_ciclo < 60:
            return "Media intensidad - Rotación moderada"
        else:
            return "Baja intensidad - Rotación lenta"

# ===============================
# 🗺️ SISTEMA DE MAPAS (interpolación KNN)
# ===============================
class SistemaMapas:
    def __init__(self):
        self.capa_base = 'https://server.arcgisonline.com/ArcGIS/rest/services/World_Imagery/MapServer/tile/{z}/{y}/{x}'
        self.estilos = {
            'area_estudio': {
                'fillColor': '#3b82f6',
                'color': '#1d4ed8',
                'weight': 4,
                'fillOpacity': 0.15,
                'dashArray': '5, 5'
            },
            'gradientes': {
                'carbono': {
                    0.0: '#0000FF', 0.2: '#00FFFF', 0.4: '#00FF00', 0.6: '#FFFF00', 0.8: '#FFA500', 1.0: '#FF0000'
                },
                'ndvi': {
                    0.0: '#8B0000', 0.2: '#FF4500', 0.4: '#FFD700', 0.6: '#9ACD32', 0.8: '#32CD32', 1.0: '#006400'
                },
                'ndwi': {
                    0.0: '#8B4513', 0.2: '#D2691E', 0.4: '#F4A460', 0.6: '#87CEEB', 0.8: '#1E90FF', 1.0: '#00008B'
                },
                'biodiversidad': {
                    0.0: '#991B1B', 0.2: '#EF4444', 0.4: '#F59E0B', 0.6: '#3B82F6', 0.8: '#8B5CF6', 1.0: '#10B981'
                },
                'forraje': {
                    0.0: '#8B4513', 0.2: '#CD853F', 0.4: '#F4A460', 0.6: '#9ACD32', 0.8: '#32CD32', 1.0: '#006400'
                },
                'ndre': {
                    0.0: '#8B0000', 0.2: '#FF4500', 0.4: '#FFD700', 0.6: '#7CFC00', 0.8: '#32CD32', 1.0: '#006400'
                },
                'msavi': {
                    0.0: '#8B4513', 0.2: '#CD853F', 0.4: '#F4A460', 0.6: '#9ACD32', 0.8: '#32CD32', 1.0: '#006400'
                },
                'evi': {
                    0.0: '#8B0000', 0.2: '#FF6347', 0.4: '#FFD700', 0.6: '#7CFC00', 0.8: '#32CD32', 1.0: '#006400'
                }
            }
        }

    def _generar_malla_puntos(self, gdf, densidad=1200):
        if gdf is None or gdf.empty:
            return []
        try:
            poligono = gdf.geometry.iloc[0]
            bounds = gdf.total_bounds
            minx, miny, maxx, maxy = bounds
            area_ha = calcular_superficie(gdf)
            num_puntos = min(densidad, max(400, int(area_ha * 1.5)))
            puntos = []
            lado = int(np.sqrt(num_puntos))
            dx = (maxx - minx) / lado
            dy = (maxy - miny) / lado
            for i in range(lado):
                for j in range(lado):
                    lon = minx + (i + 0.5) * dx
                    lat = miny + (j + 0.5) * dy
                    punto = Point(lon, lat)
                    if poligono.contains(punto):
                        puntos.append({'lat': lat, 'lon': lon, 'x_norm': i / lado, 'y_norm': j / lado})
            return puntos
        except Exception as e:
            print(f"Error generando malla: {str(e)}")
            return []

    def _interpolar_valores_knn(self, puntos_muestra, puntos_malla, variable='carbono', k=8):
        if not puntos_muestra or not puntos_malla:
            return puntos_malla
        try:
            from sklearn.neighbors import KNeighborsRegressor
            sklearn_disponible = True
        except ImportError:
            sklearn_disponible = False

        if sklearn_disponible:
            X_train = []
            y_train = []
            for punto in puntos_muestra:
                X_train.append([punto['lat'], punto['lon']])
                if variable == 'carbono':
                    y_train.append(punto['carbono_ton_ha'])
                elif variable == 'ndvi':
                    y_train.append(punto['ndvi'])
                elif variable == 'ndwi':
                    y_train.append(punto['ndwi'])
                elif variable == 'biodiversidad':
                    y_train.append(punto['indice_shannon'])
                elif variable == 'forraje':
                    y_train.append(punto['productividad_kg_ms_ha'])
                elif variable == 'ndre':
                    y_train.append(punto['ndre'])
                elif variable == 'msavi':
                    y_train.append(punto['msavi'])
                elif variable == 'evi':
                    y_train.append(punto['evi'])

            knn = KNeighborsRegressor(n_neighbors=min(k, len(X_train)), weights='distance')
            knn.fit(X_train, y_train)
            X_pred = [[p['lat'], p['lon']] for p in puntos_malla]
            if len(X_pred) > 0:
                predicciones = knn.predict(X_pred)
                for i, punto in enumerate(puntos_malla):
                    valor = float(predicciones[i])
                    if variable == 'carbono':
                        punto['carbono_ton_ha'] = max(0, valor)
                    elif variable == 'ndvi':
                        punto['ndvi'] = max(-1.0, min(1.0, valor))
                    elif variable == 'ndwi':
                        punto['ndwi'] = max(-1.0, min(1.0, valor))
                    elif variable == 'biodiversidad':
                        punto['indice_shannon'] = max(0, valor)
                    elif variable == 'forraje':
                        punto['productividad_kg_ms_ha'] = max(0, valor)
                    elif variable == 'ndre':
                        punto['ndre'] = max(-1.0, min(1.0, valor))
                    elif variable == 'msavi':
                        punto['msavi'] = max(0, valor)
                    elif variable == 'evi':
                        punto['evi'] = max(0, valor)
        else:
            for punto_malla in puntos_malla:
                valores = []
                distancias = []
                for punto_muestra in puntos_muestra:
                    dist = np.sqrt((punto_malla['lat'] - punto_muestra['lat'])**2 + (punto_malla['lon'] - punto_muestra['lon'])**2)
                    if variable == 'carbono':
                        valor = punto_muestra['carbono_ton_ha']
                    elif variable == 'ndvi':
                        valor = punto_muestra['ndvi']
                    elif variable == 'ndwi':
                        valor = punto_muestra['ndwi']
                    elif variable == 'biodiversidad':
                        valor = punto_muestra['indice_shannon']
                    elif variable == 'forraje':
                        valor = punto_muestra['productividad_kg_ms_ha']
                    elif variable == 'ndre':
                        valor = punto_muestra['ndre']
                    elif variable == 'msavi':
                        valor = punto_muestra['msavi']
                    elif variable == 'evi':
                        valor = punto_muestra['evi']
                    peso = 1.0 / (dist ** 2) if dist > 0 else 1.0
                    valores.append(valor)
                    distancias.append(peso)
                if distancias:
                    total_pesos = sum(distancias)
                    valor_interpolado = sum(v * w for v, w in zip(valores, distancias)) / total_pesos if total_pesos > 0 else np.mean(valores)
                else:
                    valor_interpolado = 0
                if variable == 'carbono':
                    punto_malla['carbono_ton_ha'] = max(0, valor_interpolado)
                elif variable == 'ndvi':
                    punto_malla['ndvi'] = max(-1.0, min(1.0, valor_interpolado))
                elif variable == 'ndwi':
                    punto_malla['ndwi'] = max(-1.0, min(1.0, valor_interpolado))
                elif variable == 'biodiversidad':
                    punto_malla['indice_shannon'] = max(0, valor_interpolado)
                elif variable == 'forraje':
                    punto_malla['productividad_kg_ms_ha'] = max(0, valor_interpolado)
                elif variable == 'ndre':
                    punto_malla['ndre'] = max(-1.0, min(1.0, valor_interpolado))
                elif variable == 'msavi':
                    punto_malla['msavi'] = max(0, valor_interpolado)
                elif variable == 'evi':
                    punto_malla['evi'] = max(0, valor_interpolado)
        return puntos_malla

    def crear_mapa_area(self, gdf, zoom_auto=True):
        if gdf is None or gdf.empty:
            return None
        try:
            bounds = gdf.total_bounds
            centro = [(bounds[1] + bounds[3]) / 2, (bounds[0] + bounds[2]) / 2]
            if zoom_auto:
                width = bounds[2] - bounds[0]
                height = bounds[3] - bounds[1]
                extension = max(width, height)
                if extension > 10: zoom_start = 6
                elif extension > 5: zoom_start = 8
                elif extension > 2: zoom_start = 10
                elif extension > 1: zoom_start = 12
                elif extension > 0.5: zoom_start = 14
                elif extension > 0.2: zoom_start = 16
                else: zoom_start = 18
            else:
                zoom_start = 12
            m = folium.Map(location=centro, zoom_start=zoom_start, tiles=self.capa_base, attr='Esri, Maxar, Earthstar Geographics', control_scale=True)
            folium.GeoJson(gdf.geometry.iloc[0], style_function=lambda x: self.estilos['area_estudio'],
                           highlight_function=lambda x: {'weight': 6, 'color': '#1e40af', 'fillOpacity': 0.3}).add_to(m)
            m.fit_bounds([[bounds[1], bounds[0]], [bounds[3], bounds[2]]])
            Fullscreen().add_to(m)
            MousePosition().add_to(m)
            return m
        except Exception as e:
            st.warning(f"Error al crear mapa: {str(e)}")
            return None

    def crear_mapa_calor_interpolado(self, resultados, variable='carbono', gdf_area=None):
        if not resultados or gdf_area is None or gdf_area.empty:
            return None
        try:
            puntos_muestra = resultados.get(f'puntos_{variable}', [])
            if not puntos_muestra:
                return None
            puntos_malla = self._generar_malla_puntos(gdf_area, densidad=1200)
            if not puntos_malla:
                return None
            puntos_interpolados = self._interpolar_valores_knn(puntos_muestra, puntos_malla, variable)
            bounds = gdf_area.total_bounds
            centro = [(bounds[1] + bounds[3]) / 2, (bounds[0] + bounds[2]) / 2]
            m = folium.Map(location=centro, zoom_start=12, tiles=self.capa_base, attr='Esri, Maxar, Earthstar Geographics', control_scale=True)
            folium.GeoJson(gdf_area.geometry.iloc[0], style_function=lambda x: {
                'fillColor': 'transparent', 'color': '#1d4ed8', 'weight': 2, 'fillOpacity': 0.05, 'dashArray': '5, 5'
            }).add_to(m)
            heat_data = []
            for punto in puntos_interpolados:
                if variable == 'carbono':
                    heat_data.append([punto['lat'], punto['lon'], punto['carbono_ton_ha']])
                elif variable == 'ndvi':
                    heat_data.append([punto['lat'], punto['lon'], punto['ndvi']])
                elif variable == 'ndwi':
                    heat_data.append([punto['lat'], punto['lon'], punto['ndwi']])
                elif variable == 'biodiversidad':
                    heat_data.append([punto['lat'], punto['lon'], punto['indice_shannon']])
                elif variable == 'forraje':
                    heat_data.append([punto['lat'], punto['lon'], punto['productividad_kg_ms_ha']])
                elif variable == 'ndre':
                    heat_data.append([punto['lat'], punto['lon'], punto['ndre']])
                elif variable == 'msavi':
                    heat_data.append([punto['lat'], punto['lon'], punto['msavi']])
                elif variable == 'evi':
                    heat_data.append([punto['lat'], punto['lon'], punto['evi']])
            gradient = self.estilos['gradientes'].get(variable, self.estilos['gradientes']['carbono'])
            radius = 45 if variable in ['carbono', 'biodiversidad', 'forraje'] else 40
            blur = 40 if variable in ['carbono', 'biodiversidad', 'forraje'] else 35
            HeatMap(heat_data, name=variable, min_opacity=0.7, radius=radius, blur=blur, gradient=gradient, max_zoom=18).add_to(m)
            m.fit_bounds([[bounds[1], bounds[0]], [bounds[3], bounds[2]]])
            return m
        except Exception as e:
            st.warning(f"Error al crear mapa de calor para {variable}: {str(e)}")
            return None

    def crear_mapa_estatico(self, resultados, variable='carbono', gdf_area=None, dpi=150):
        if not resultados or gdf_area is None or gdf_area.empty:
            return None
        puntos_muestra = resultados.get(f'puntos_{variable}', [])
        if not puntos_muestra:
            return None
        puntos_malla = self._generar_malla_puntos(gdf_area, densidad=800)
        if not puntos_malla:
            return None
        puntos_interpolados = self._interpolar_valores_knn(puntos_muestra, puntos_malla, variable)
        lats = [p['lat'] for p in puntos_interpolados]
        lons = [p['lon'] for p in puntos_interpolados]
        if variable == 'carbono':
            valores = [p['carbono_ton_ha'] for p in puntos_interpolados]
            titulo = 'Carbono (ton C/ha)'
            cmap_name = 'carbono'
        elif variable == 'ndvi':
            valores = [p['ndvi'] for p in puntos_interpolados]
            titulo = 'NDVI'
            cmap_name = 'ndvi'
        elif variable == 'ndwi':
            valores = [p['ndwi'] for p in puntos_interpolados]
            titulo = 'NDWI'
            cmap_name = 'ndwi'
        elif variable == 'biodiversidad':
            valores = [p['indice_shannon'] for p in puntos_interpolados]
            titulo = 'Índice de Shannon'
            cmap_name = 'biodiversidad'
        elif variable == 'forraje':
            valores = [p['productividad_kg_ms_ha'] for p in puntos_interpolados]
            titulo = 'Productividad (kg MS/ha)'
            cmap_name = 'forraje'
        elif variable == 'ndre':
            valores = [p['ndre'] for p in puntos_interpolados]
            titulo = 'NDRE'
            cmap_name = 'ndre'
        elif variable == 'msavi':
            valores = [p['msavi'] for p in puntos_interpolados]
            titulo = 'MSAVI'
            cmap_name = 'msavi'
        elif variable == 'evi':
            valores = [p['evi'] for p in puntos_interpolados]
            titulo = 'EVI'
            cmap_name = 'evi'
        else:
            return None
        bounds = gdf_area.total_bounds
        minx, miny, maxx, maxy = bounds
        grid_x, grid_y = np.mgrid[minx:maxx:100j, miny:maxy:100j]
        grid_z = griddata((lons, lats), valores, (grid_x, grid_y), method='cubic')
        fig, ax = plt.subplots(1, 1, figsize=(10, 8))
        colormap = LinearSegmentedColormap.from_list(cmap_name, list(self.estilos['gradientes'][cmap_name].values()))
        im = ax.imshow(grid_z.T, extent=[minx, maxx, miny, maxy], origin='lower', cmap=colormap, aspect='auto')
        plt.colorbar(im, ax=ax, label=titulo)
        ax.set_title(f'Mapa de {titulo}')
        ax.set_xlabel('Longitud')
        ax.set_ylabel('Latitud')
        ax.grid(True, linestyle='--', alpha=0.5)
        if gdf_area is not None and not gdf_area.empty:
            boundary_geom = gdf_area.geometry.iloc[0].boundary
            if boundary_geom and not boundary_geom.is_empty:
                gpd.GeoSeries([boundary_geom]).plot(ax=ax, color='black', linewidth=1.5)
        buf = BytesIO()  # CORREGIDO: io.BytesIO() -> BytesIO()
        plt.savefig(buf, format='png', dpi=dpi, bbox_inches='tight')
        plt.close(fig)
        buf.seek(0)
        return buf

# ===============================
# 📊 VISUALIZACIONES
# ===============================
class Visualizaciones:
    @staticmethod
    def crear_grafico_barras_carbono(desglose: Dict):
        if not desglose:
            fig = go.Figure()
            fig.update_layout(title='No hay datos de carbono disponibles', height=400)
            return fig
        descripciones = {'AGB': 'Biomasa Aérea Viva', 'BGB': 'Biomasa de Raíces', 'DW': 'Madera Muerta', 'LI': 'Hojarasca', 'SOC': 'Carbono Orgánico del Suelo'}
        etiquetas = [f"{descripciones.get(k, k)}<br>({k})" for k in desglose.keys()]
        fig = go.Figure(data=[go.Bar(x=etiquetas, y=list(desglose.values()), marker_color=['#238b45', '#41ab5d', '#74c476', '#a1d99b', '#d9f0a3'], text=[f"{v:.1f} ton C/ha" for v in desglose.values()], textposition='auto', hovertemplate='<b>%{x}</b><br>Valor: %{y:.1f} ton C/ha<extra></extra>')])
        fig.update_layout(title='Distribución de Carbono por Pools', xaxis_title='Pool de Carbono', yaxis_title='Ton C/ha', height=400, hovermode='x unified')
        return fig

    @staticmethod
    def crear_grafico_radar_biodiversidad(shannon_data: Dict):
        if not shannon_data:
            fig = go.Figure()
            fig.update_layout(title='No hay datos de biodiversidad disponibles', height=400)
            return fig
        categorias = ['Shannon', 'Riqueza', 'Abundancia', 'Equitatividad', 'Conservación']
        try:
            shannon_norm = min(shannon_data.get('indice_shannon', 0) / 4.0 * 100, 100)
            riqueza_norm = min(shannon_data.get('riqueza_especies', 0) / 200 * 100, 100)
            abundancia_norm = min(shannon_data.get('abundancia_total', 0) / 2000 * 100, 100)
            equitatividad = random.uniform(70, 90)
            conservacion = random.uniform(60, 95)
            valores = [shannon_norm, riqueza_norm, abundancia_norm, equitatividad, conservacion]
            fig = go.Figure(data=go.Scatterpolar(r=valores, theta=categorias, fill='toself', fillcolor='rgba(139, 92, 246, 0.3)', line_color='#8b5cf6', name='Biodiversidad'))
            fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])), showlegend=True, height=400, title='Perfil de Biodiversidad')
            return fig
        except Exception as e:
            fig = go.Figure()
            fig.update_layout(title='Error al generar gráfico de biodiversidad', height=400)
            return fig

    @staticmethod
    def crear_grafico_comparativo(puntos_carbono, puntos_ndvi, puntos_ndwi, puntos_biodiversidad):
        if not puntos_carbono or not puntos_ndvi:
            return None
        try:
            n = min(50, len(puntos_carbono))
            fig = make_subplots(rows=2, cols=2, subplot_titles=('Carbono vs NDVI', 'Carbono vs NDWI', 'Shannon vs NDVI', 'Shannon vs NDWI'), vertical_spacing=0.15, horizontal_spacing=0.15)
            carbono_vals = [p['carbono_ton_ha'] for p in puntos_carbono[:n]]
            ndvi_vals = [p['ndvi'] for p in puntos_ndvi[:n]]
            ndwi_vals = [p['ndwi'] for p in puntos_ndwi[:n]]
            shannon_vals = [p['indice_shannon'] for p in puntos_biodiversidad[:n]]
            fig.add_trace(go.Scatter(x=ndvi_vals, y=carbono_vals, mode='markers', marker=dict(color='#10b981', size=8), name='Carbono-NDVI'), row=1, col=1)
            fig.add_trace(go.Scatter(x=ndwi_vals, y=carbono_vals, mode='markers', marker=dict(color='#3b82f6', size=8), name='Carbono-NDWI'), row=1, col=2)
            fig.add_trace(go.Scatter(x=ndvi_vals, y=shannon_vals, mode='markers', marker=dict(color='#8b5cf6', size=8), name='Shannon-NDVI'), row=2, col=1)
            fig.add_trace(go.Scatter(x=ndwi_vals, y=shannon_vals, mode='markers', marker=dict(color='#f59e0b', size=8), name='Shannon-NDWI'), row=2, col=2)
            fig.update_layout(height=700, showlegend=True, title_text="Comparación de Variables Ambientales")
            fig.update_xaxes(title_text="NDVI", row=1, col=1)
            fig.update_yaxes(title_text="Carbono (ton C/ha)", row=1, col=1)
            fig.update_xaxes(title_text="NDWI", row=1, col=2)
            fig.update_yaxes(title_text="Carbono (ton C/ha)", row=1, col=2)
            fig.update_xaxes(title_text="NDVI", row=2, col=1)
            fig.update_yaxes(title_text="Índice de Shannon", row=2, col=1)
            fig.update_xaxes(title_text="NDWI", row=2, col=2)
            fig.update_yaxes(title_text="Índice de Shannon", row=2, col=2)
            return fig
        except Exception as e:
            return None

    @staticmethod
    def crear_grafico_forrajero(disponibilidad_forrajera: Dict, equivalentes_vaca: Dict):
        fig = make_subplots(rows=2, cols=2, subplot_titles=('Disponibilidad Forrajera', 'Equivalentes Vaca', 'Distribución por Sublote', 'Plan de Rotación'),
                             specs=[[{'type': 'bar'}, {'type': 'pie'}], [{'type': 'bar'}, {'type': 'table'}]],
                             vertical_spacing=0.15, horizontal_spacing=0.15, row_heights=[0.5, 0.5])
        fig.add_trace(go.Bar(x=['Productividad', 'Disponible Total', 'Aprovechable'],
                              y=[disponibilidad_forrajera.get('productividad_kg_ms_ha', 0),
                                 disponibilidad_forrajera.get('disponibilidad_total_kg_ms', 0) / 1000,
                                 disponibilidad_forrajera.get('forraje_aprovechable_kg_ms', 0) / 1000],
                              name='Forraje', marker_color=['#8B4513', '#D2691E', '#F4A460']), row=1, col=1)
        fig.add_trace(go.Pie(labels=['EV por día', 'EV para período', 'EV recomendado'],
                              values=[equivalentes_vaca.get('ev_por_dia', 0),
                                      equivalentes_vaca.get('ev_para_periodo', 0),
                                      equivalentes_vaca.get('ev_recomendado', 0)],
                              name='Equivalentes Vaca', hole=0.4), row=1, col=2)
        fig.update_layout(height=700, showlegend=True, title_text="Análisis Forrajero Completo")
        fig.update_yaxes(title_text="kg MS/ha / ton MS", row=1, col=1)
        fig.update_xaxes(title_text="Métrica", row=1, col=1)
        return fig

    @staticmethod
    def crear_metricas_kpi(carbono_total: float, co2_total: float, shannon: float, area: float):
        html = f"""
        <div style="display: grid; grid-template-columns: repeat(4, 1fr); gap: 1rem; margin-bottom: 2rem;">
            <div style="background: linear-gradient(135deg, #065f46 0%, #0a7e5a 100%); padding: 1.5rem; border-radius: 10px; color: white;">
                <h3 style="margin: 0; font-size: 1.2rem;">🌳 Carbono Total</h3>
                <p style="font-size: 2rem; font-weight: bold; margin: 0.5rem 0;">{carbono_total:,.0f}</p>
                <p style="margin: 0;">ton C</p>
            </div>
            <div style="background: linear-gradient(135deg, #0a7e5a 0%, #10b981 100%); padding: 1.5rem; border-radius: 10px; color: white;">
                <h3 style="margin: 0; font-size: 1.2rem;">🏭 CO₂ Equivalente</h3>
                <p style="font-size: 2rem; font-weight: bold; margin: 0.5rem 0;">{co2_total:,.0f}</p>
                <p style="margin: 0;">ton CO₂e</p>
            </div>
            <div style="background: linear-gradient(135deg, #8b5cf6 0%, #a78bfa 100%); padding: 1.5rem; border-radius: 10px; color: white;">
                <h3 style="margin: 0; font-size: 1.2rem;">🦋 Índice Shannon</h3>
                <p style="font-size: 2rem; font-weight: bold; margin: 0.5rem 0;">{shannon:.2f}</p>
                <p style="margin: 0;">Biodiversidad</p>
            </div>
            <div style="background: linear-gradient(135deg, #3b82f6 0%, #60a5fa 100%); padding: 1.5rem; border-radius: 10px; color: white;">
                <h3 style="margin: 0; font-size: 1.2rem;">📐 Área Total</h3>
                <p style="font-size: 2rem; font-weight: bold; margin: 0.5rem 0;">{area:,.1f}</p>
                <p style="margin: 0;">hectáreas</p>
            </div>
        </div>
        """
        return html

# ===============================
# 📄 GENERADOR DE REPORTES
# ===============================
class GeneradorReportes:
    def __init__(self, resultados, gdf, sistema_mapas=None):
        self.resultados = resultados
        self.gdf = gdf
        self.sistema_mapas = sistema_mapas
        self.buffer_pdf = BytesIO()
        self.buffer_docx = BytesIO()

    def _fig_to_png(self, fig, width=800, height=500):
        if fig is None:
            return None
        try:
            img_bytes = fig.to_image(format='png', width=width, height=height, scale=2)
            return BytesIO(img_bytes)
        except Exception as e:
            st.warning(f"No se pudo convertir el gráfico a PNG: {str(e)}")
            return None

    def _mapa_to_png(self, mapa, width=800, height=600):
        try:
            if mapa is None:
                return None
            from PIL import Image, ImageDraw
            img = Image.new('RGB', (width, height), color='white')
            draw = ImageDraw.Draw(img)
            draw.text((width//2 - 100, height//2 - 20), "Mapa interactivo", fill='black')
            draw.text((width//2 - 150, height//2 + 10), "Disponible en la aplicación web", fill='gray')
            draw.rectangle([10, 10, width-10, height-10], outline='blue', width=3)
            img_byte_arr = BytesIO()  # CORREGIDO: io.BytesIO() -> BytesIO()
            img.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            return img_byte_arr
        except Exception as e:
            st.warning(f"No se pudo convertir el mapa a PNG: {str(e)}")
            return None

    def _crear_graficos(self):
        vis = Visualizaciones()
        res = self.resultados
        graficos = {}
        if 'desglose_promedio' in res and res['desglose_promedio']:
            fig_carbono = vis.crear_grafico_barras_carbono(res['desglose_promedio'])
            graficos['carbono'] = self._fig_to_png(fig_carbono)
        if 'puntos_biodiversidad' in res and res['puntos_biodiversidad'] and len(res['puntos_biodiversidad']) > 0:
            fig_biodiv = vis.crear_grafico_radar_biodiversidad(res['puntos_biodiversidad'][0])
            graficos['biodiv'] = self._fig_to_png(fig_biodiv)
        if all(k in res for k in ['puntos_carbono', 'puntos_ndvi', 'puntos_ndwi', 'puntos_biodiversidad']):
            fig_comparativo = vis.crear_grafico_comparativo(
                res['puntos_carbono'], res['puntos_ndvi'], res['puntos_ndwi'], res['puntos_biodiversidad']
            )
            if fig_comparativo:
                graficos['comparativo'] = self._fig_to_png(fig_comparativo)
        if 'analisis_forrajero' in res:
            forrajero_data = res['analisis_forrajero']
            if 'disponibilidad_forrajera' in forrajero_data and 'equivalentes_vaca' in forrajero_data:
                fig_forrajero = vis.crear_grafico_forrajero(
                    forrajero_data['disponibilidad_forrajera'],
                    forrajero_data['equivalentes_vaca']
                )
                graficos['forrajero'] = self._fig_to_png(fig_forrajero)
        return graficos

    def generar_pdf(self):
        if not REPORTPDF_AVAILABLE:
            st.error("ReportLab no está instalado. No se puede generar PDF.")
            return None
        try:
            doc = SimpleDocTemplate(self.buffer_pdf, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
            story = []
            styles = getSampleStyleSheet()
            titulo_style = ParagraphStyle('TituloPrincipal', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#0a7e5a'), spaceAfter=30, alignment=TA_CENTER)
            subtitulo_style = ParagraphStyle('Subtitulo', parent=styles['Heading2'], fontSize=18, textColor=colors.HexColor('#065f46'), spaceAfter=12, spaceBefore=20)
            seccion_style = ParagraphStyle('Seccion', parent=styles['Heading3'], fontSize=14, textColor=colors.HexColor('#1d4ed8'), spaceAfter=10, spaceBefore=15)
            # Portada
            story.append(Paragraph("INFORME AMBIENTAL INTEGRAL", titulo_style))
            story.append(Spacer(1, 12))
            story.append(Paragraph("Sistema Satelital de Análisis Ambiental", styles['Title']))
            story.append(Spacer(1, 6))
            story.append(Paragraph("Carbono + Biodiversidad + Análisis Forrajero", styles['Heading2']))
            story.append(Spacer(1, 24))
            story.append(Paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles['Normal']))
            story.append(Spacer(1, 36))
            # Resumen ejecutivo
            story.append(Paragraph("RESUMEN EJECUTIVO", subtitulo_style))
            res = self.resultados
            datos_resumen = [
                ["Métrica", "Valor", "Interpretación"],
                ["Área total", f"{res.get('area_total_ha', 0):,.1f} ha", "Superficie del área de estudio"],
                ["Carbono total almacenado", f"{res.get('carbono_total_ton', 0):,.0f} ton C", "Carbono almacenado en el área"],
                ["CO₂ equivalente", f"{res.get('co2_total_ton', 0):,.0f} ton CO₂e", "Potencial de créditos de carbono"],
                ["Índice de Shannon promedio", f"{res.get('shannon_promedio', 0):.3f}", "Nivel de biodiversidad"],
                ["NDVI promedio", f"{res.get('ndvi_promedio', 0):.3f}", "Salud de la vegetación"],
                ["NDWI promedio", f"{res.get('ndwi_promedio', 0):.3f}", "Contenido de agua"],
                ["Tipo de ecosistema", res.get('tipo_ecosistema', 'N/A'), "Ecosistema predominante"],
                ["Puntos de muestreo", str(res.get('num_puntos', 0)), "Muestras analizadas"]
            ]
            tabla_resumen = Table(datos_resumen, colWidths=[150, 120, 200])
            tabla_resumen.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#065f46')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f0f9ff')),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e1')),
                ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            story.append(tabla_resumen)
            story.append(Spacer(1, 20))
            # Análisis de carbono
            story.append(PageBreak())
            story.append(Paragraph("ANÁLISIS DE CARBONO", subtitulo_style))
            if res.get('desglose_promedio'):
                descripciones = {
                    'AGB': 'Biomasa Aérea Viva', 'BGB': 'Biomasa de Raíces', 'DW': 'Madera Muerta',
                    'LI': 'Hojarasca', 'SOC': 'Carbono Orgánico del Suelo'
                }
                datos_carbono = [["Pool", "Descripción", "Ton C/ha", "Porcentaje"]]
                total = sum(res['desglose_promedio'].values())
                for pool, valor in res['desglose_promedio'].items():
                    porcentaje = (valor / total * 100) if total > 0 else 0
                    datos_carbono.append([pool, descripciones.get(pool, pool), f"{valor:.2f}", f"{porcentaje:.1f}%"])
                tabla_carbono = Table(datos_carbono, colWidths=[60, 180, 70, 70])
                tabla_carbono.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0a7e5a')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('ALIGN', (2, 1), (3, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f0fdf4')),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#d1fae5')),
                ]))
                story.append(tabla_carbono)
                story.append(Spacer(1, 15))
            # Análisis de biodiversidad
            story.append(PageBreak())
            story.append(Paragraph("ANÁLISIS DE BIODIVERSIDAD", subtitulo_style))
            if res.get('puntos_biodiversidad') and len(res['puntos_biodiversidad']) > 0:
                biodiv = res['puntos_biodiversidad'][0]
                datos_biodiv = [
                    ["Métrica", "Valor", "Interpretación"],
                    ["Índice de Shannon", f"{biodiv.get('indice_shannon', 0):.3f}", biodiv.get('categoria', 'N/A')],
                    ["Riqueza de especies", str(biodiv.get('riqueza_especies', 0)), "Número estimado de especies"],
                    ["Abundancia total", f"{biodiv.get('abundancia_total', 0):,}", "Individuos estimados"],
                    ["Categoría", biodiv.get('categoria', 'N/A'), "Clasificación según Shannon"]
                ]
                tabla_biodiv = Table(datos_biodiv, colWidths=[120, 100, 180])
                tabla_biodiv.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8b5cf6')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                    ('ALIGN', (1, 1), (1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#faf5ff')),
                    ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e9d5ff')),
                ]))
                story.append(tabla_biodiv)
                story.append(Spacer(1, 15))
            # Análisis forrajero
            story.append(PageBreak())
            story.append(Paragraph("ANÁLISIS FORRAJERO", subtitulo_style))
            if 'analisis_forrajero' in res:
                forrajero_data = res['analisis_forrajero']
                if 'disponibilidad_forrajera' in forrajero_data:
                    disp = forrajero_data['disponibilidad_forrajera']
                    datos_forraje = [
                        ["Métrica", "Valor", "Unidad"],
                        ["Productividad", f"{disp.get('productividad_kg_ms_ha', 0):,.0f}", "kg MS/ha"],
                        ["Disponibilidad total", f"{disp.get('disponibilidad_total_kg_ms', 0)/1000:,.1f}", "ton MS"],
                        ["Forraje aprovechable", f"{disp.get('forraje_aprovechable_kg_ms', 0)/1000:,.1f}", "ton MS"],
                        ["Tasa crecimiento diario", f"{disp.get('tasa_crecimiento_diario_kg', 0):,.0f}", "kg/día"],
                        ["Categoría productividad", disp.get('categoria_productividad', 'N/A').title(), ""]
                    ]
                    tabla_forraje = Table(datos_forraje, colWidths=[150, 100, 80])
                    tabla_forraje.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8B4513')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('ALIGN', (1, 1), (2, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#fdf4e3')),
                        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#d2b48c')),
                    ]))
                    story.append(tabla_forraje)
                    story.append(Spacer(1, 15))
                if 'equivalentes_vaca' in forrajero_data:
                    ev = forrajero_data['equivalentes_vaca']
                    datos_ev = [
                        ["Concepto", "Valor"],
                        ["EV por día", f"{ev.get('ev_por_dia', 0):.1f}"],
                        ["EV para 30 días", f"{ev.get('ev_para_periodo', 0):.1f}"],
                        ["EV recomendado", f"{ev.get('ev_recomendado', 0):.1f}"],
                        ["Consumo EV diario", f"{ev.get('consumo_ev_diario_kg', 0)} kg"]
                    ]
                    tabla_ev = Table(datos_ev, colWidths=[150, 100])
                    tabla_ev.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#CD853F')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#fff8dc')),
                        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cd853f')),
                    ]))
                    story.append(tabla_ev)
                    story.append(Spacer(1, 15))
                if 'sublotes' in forrajero_data and forrajero_data['sublotes']:
                    datos_sublotes = [["Sublote", "Área (ha)", "Productividad (kg MS/ha)", "Forraje aprovechable (ton)"]]
                    for s in forrajero_data['sublotes']:
                        datos_sublotes.append([
                            str(s['sublote_id']),
                            f"{s['area_ha']:.1f}",
                            f"{s['disponibilidad_kg_ms_ha']:,.0f}",
                            f"{s['forraje_aprovechable_kg_ms']/1000:.1f}"
                        ])
                    tabla_sublotes = Table(datos_sublotes, colWidths=[60, 70, 120, 100])
                    tabla_sublotes.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#8B4513')),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                        ('ALIGN', (1, 1), (3, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#fdf4e3')),
                        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#d2b48c')),
                    ]))
                    story.append(tabla_sublotes)
                    story.append(Spacer(1, 15))
            # Índices espectrales
            story.append(PageBreak())
            story.append(Paragraph("ÍNDICES ESPECTRALES", subtitulo_style))
            datos_indices = [
                ["Índice", "Valor promedio"],
                ["NDVI", f"{res.get('ndvi_promedio', 0):.3f}"],
                ["NDWI", f"{res.get('ndwi_promedio', 0):.3f}"],
            ]
            if 'puntos_ndre' in res:
                ndre_vals = [p['ndre'] for p in res['puntos_ndre']]
                datos_indices.append(["NDRE", f"{np.mean(ndre_vals):.3f}"])
            if 'puntos_msavi' in res:
                msavi_vals = [p['msavi'] for p in res['puntos_msavi']]
                datos_indices.append(["MSAVI", f"{np.mean(msavi_vals):.3f}"])
            if 'puntos_evi' in res:
                evi_vals = [p['evi'] for p in res['puntos_evi']]
                datos_indices.append(["EVI", f"{np.mean(evi_vals):.3f}"])
            tabla_indices = Table(datos_indices, colWidths=[100, 100])
            tabla_indices.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#10b981')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f0fdf4')),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#bbf7d0')),
            ]))
            story.append(tabla_indices)
            story.append(Spacer(1, 20))
            # Mapas estáticos (si los hay)
            if self.sistema_mapas:
                story.append(PageBreak())
                story.append(Paragraph("MAPAS DE CALOR", subtitulo_style))
                variables = ['carbono', 'ndvi', 'ndwi', 'biodiversidad', 'forraje']
                for var in variables:
                    mapa = self.sistema_mapas.crear_mapa_estatico(self.resultados, var, self.gdf)
                    if mapa:
                        story.append(Paragraph(f"Mapa de {var.replace('_',' ').title()}", seccion_style))
                        story.append(Image(mapa, width=450, height=350))
                        story.append(Spacer(1, 12))
            # Conclusiones
            story.append(PageBreak())
            story.append(Paragraph("CONCLUSIONES Y RECOMENDACIONES", subtitulo_style))
            if 'analisis_forrajero' in res:
                forrajero_data = res['analisis_forrajero']
            else:
                forrajero_data = {}
            conclusiones = [
                f"El área de estudio de {res.get('area_total_ha', 0):,.1f} hectáreas almacena {res.get('carbono_total_ton', 0):,.0f} ton C, equivalente a {res.get('co2_total_ton', 0):,.0f} ton CO₂e.",
                f"El índice de Shannon promedio es {res.get('shannon_promedio', 0):.3f}, lo que indica una biodiversidad {res.get('puntos_biodiversidad', [{}])[0].get('categoria', 'N/A').lower()}.",
                f"El NDVI promedio de {res.get('ndvi_promedio', 0):.3f} sugiere una cobertura vegetal moderada.",
                f"La productividad forrajera estimada es de {forrajero_data.get('disponibilidad_forrajera', {}).get('productividad_kg_ms_ha', 0):,.0f} kg MS/ha, lo que permite recomendar una carga de {forrajero_data.get('equivalentes_vaca', {}).get('ev_recomendado', 0):.1f} EV para un período de 30 días."
            ]
            for conc in conclusiones:
                story.append(Paragraph(conc, styles['Normal']))
                story.append(Spacer(1, 8))
            doc.build(story)
            self.buffer_pdf.seek(0)
            return self.buffer_pdf
        except Exception as e:
            st.error(f"Error generando PDF: {str(e)}")
            import traceback
            st.error(traceback.format_exc())
            return None

    def generar_docx(self):
        if not REPORTDOCX_AVAILABLE:
            st.error("python-docx no está instalado. No se puede generar DOCX.")
            return None
        try:
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(11)
            title = doc.add_heading('INFORME AMBIENTAL INTEGRAL', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph()
            # Resumen ejecutivo
            doc.add_heading('RESUMEN EJECUTIVO', level=1)
            res = self.resultados
            tabla_resumen = doc.add_table(rows=9, cols=3)
            tabla_resumen.style = 'Light Shading'
            tabla_resumen.cell(0, 0).text = 'Métrica'
            tabla_resumen.cell(0, 1).text = 'Valor'
            tabla_resumen.cell(0, 2).text = 'Interpretación'
            datos = [
                ('Área total', f"{res.get('area_total_ha', 0):,.1f} ha", 'Superficie del área de estudio'),
                ('Carbono total almacenado', f"{res.get('carbono_total_ton', 0):,.0f} ton C", 'Carbono almacenado en el área'),
                ('CO₂ equivalente', f"{res.get('co2_total_ton', 0):,.0f} ton CO₂e", 'Potencial de créditos de carbono'),
                ('Índice de Shannon promedio', f"{res.get('shannon_promedio', 0):.3f}", 'Nivel de biodiversidad'),
                ('NDVI promedio', f"{res.get('ndvi_promedio', 0):.3f}", 'Salud de la vegetación'),
                ('NDWI promedio', f"{res.get('ndwi_promedio', 0):.3f}", 'Contenido de agua'),
                ('Tipo de ecosistema', res.get('tipo_ecosistema', 'N/A'), 'Ecosistema predominante'),
                ('Puntos de muestreo', str(res.get('num_puntos', 0)), 'Muestras analizadas')
            ]
            for i, (met, val, interp) in enumerate(datos, 1):
                tabla_resumen.cell(i, 0).text = met
                tabla_resumen.cell(i, 1).text = val
                tabla_resumen.cell(i, 2).text = interp
            doc.add_paragraph()
            # Análisis de carbono
            doc.add_heading('ANÁLISIS DE CARBONO', level=1)
            if res.get('desglose_promedio'):
                doc.add_heading('Distribución por Pools', level=2)
                tabla_carbono = doc.add_table(rows=6, cols=4)
                tabla_carbono.style = 'Light Shading'
                tabla_carbono.cell(0, 0).text = 'Pool'
                tabla_carbono.cell(0, 1).text = 'Descripción'
                tabla_carbono.cell(0, 2).text = 'Ton C/ha'
                tabla_carbono.cell(0, 3).text = 'Porcentaje'
                desc = {'AGB': 'Biomasa Aérea Viva', 'BGB': 'Biomasa de Raíces', 'DW': 'Madera Muerta', 'LI': 'Hojarasca', 'SOC': 'Carbono Orgánico del Suelo'}
                total = sum(res['desglose_promedio'].values())
                for i, (pool, valor) in enumerate(res['desglose_promedio'].items(), 1):
                    tabla_carbono.cell(i, 0).text = pool
                    tabla_carbono.cell(i, 1).text = desc.get(pool, pool)
                    tabla_carbono.cell(i, 2).text = f"{valor:.2f}"
                    porcentaje = (valor / total * 100) if total > 0 else 0
                    tabla_carbono.cell(i, 3).text = f"{porcentaje:.1f}%"
            doc.add_page_break()
            # Análisis de biodiversidad
            doc.add_heading('ANÁLISIS DE BIODIVERSIDAD', level=1)
            if res.get('puntos_biodiversidad') and len(res['puntos_biodiversidad']) > 0:
                biodiv = res['puntos_biodiversidad'][0]
                tabla_biodiv = doc.add_table(rows=5, cols=3)
                tabla_biodiv.style = 'Light Shading'
                tabla_biodiv.cell(0, 0).text = 'Métrica'
                tabla_biodiv.cell(0, 1).text = 'Valor'
                tabla_biodiv.cell(0, 2).text = 'Interpretación'
                datos_biodiv = [
                    ('Índice de Shannon', f"{biodiv.get('indice_shannon', 0):.3f}", biodiv.get('categoria', 'N/A')),
                    ('Riqueza de especies', str(biodiv.get('riqueza_especies', 0)), 'Número estimado de especies'),
                    ('Abundancia total', f"{biodiv.get('abundancia_total', 0):,}", 'Individuos estimados'),
                    ('Categoría', biodiv.get('categoria', 'N/A'), 'Clasificación según Shannon')
                ]
                for i, (met, val, interp) in enumerate(datos_biodiv, 1):
                    tabla_biodiv.cell(i, 0).text = met
                    tabla_biodiv.cell(i, 1).text = val
                    tabla_biodiv.cell(i, 2).text = interp
            doc.add_page_break()
            # Análisis forrajero
            doc.add_heading('ANÁLISIS FORRAJERO', level=1)
            if 'analisis_forrajero' in res:
                forrajero_data = res['analisis_forrajero']
                if 'disponibilidad_forrajera' in forrajero_data:
                    disp = forrajero_data['disponibilidad_forrajera']
                    doc.add_heading('Disponibilidad Forrajera', level=2)
                    tabla_forraje = doc.add_table(rows=6, cols=3)
                    tabla_forraje.style = 'Light Shading'
                    tabla_forraje.cell(0, 0).text = 'Métrica'
                    tabla_forraje.cell(0, 1).text = 'Valor'
                    tabla_forraje.cell(0, 2).text = 'Unidad'
                    datos_f = [
                        ('Productividad', f"{disp.get('productividad_kg_ms_ha', 0):,.0f}", 'kg MS/ha'),
                        ('Disponibilidad total', f"{disp.get('disponibilidad_total_kg_ms', 0)/1000:,.1f}", 'ton MS'),
                        ('Forraje aprovechable', f"{disp.get('forraje_aprovechable_kg_ms', 0)/1000:,.1f}", 'ton MS'),
                        ('Tasa crecimiento diario', f"{disp.get('tasa_crecimiento_diario_kg', 0):,.0f}", 'kg/día'),
                        ('Categoría productividad', disp.get('categoria_productividad', 'N/A').title(), '')
                    ]
                    for i, (met, val, uni) in enumerate(datos_f, 1):
                        tabla_forraje.cell(i, 0).text = met
                        tabla_forraje.cell(i, 1).text = val
                        tabla_forraje.cell(i, 2).text = uni
                if 'equivalentes_vaca' in forrajero_data:
                    ev = forrajero_data['equivalentes_vaca']
                    doc.add_heading('Equivalentes Vaca', level=2)
                    tabla_ev = doc.add_table(rows=5, cols=2)
                    tabla_ev.style = 'Light Shading'
                    tabla_ev.cell(0, 0).text = 'Concepto'
                    tabla_ev.cell(0, 1).text = 'Valor'
                    datos_ev = [
                        ('EV por día', f"{ev.get('ev_por_dia', 0):.1f}"),
                        ('EV para 30 días', f"{ev.get('ev_para_periodo', 0):.1f}"),
                        ('EV recomendado', f"{ev.get('ev_recomendado', 0):.1f}"),
                        ('Consumo EV diario', f"{ev.get('consumo_ev_diario_kg', 0)} kg")
                    ]
                    for i, (concepto, valor) in enumerate(datos_ev, 1):
                        tabla_ev.cell(i, 0).text = concepto
                        tabla_ev.cell(i, 1).text = valor
                if 'sublotes' in forrajero_data and forrajero_data['sublotes']:
                    doc.add_heading('Sublotes', level=2)
                    tabla_sub = doc.add_table(rows=len(forrajero_data['sublotes'])+1, cols=4)
                    tabla_sub.style = 'Light Shading'
                    tabla_sub.cell(0, 0).text = 'Sublote'
                    tabla_sub.cell(0, 1).text = 'Área (ha)'
                    tabla_sub.cell(0, 2).text = 'Productividad (kg MS/ha)'
                    tabla_sub.cell(0, 3).text = 'Forraje aprovechable (ton)'
                    for i, s in enumerate(forrajero_data['sublotes'], 1):
                        tabla_sub.cell(i, 0).text = str(s['sublote_id'])
                        tabla_sub.cell(i, 1).text = f"{s['area_ha']:.1f}"
                        tabla_sub.cell(i, 2).text = f"{s['disponibilidad_kg_ms_ha']:,.0f}"
                        tabla_sub.cell(i, 3).text = f"{s['forraje_aprovechable_kg_ms']/1000:.1f}"
            doc.save(self.buffer_docx)
            self.buffer_docx.seek(0)
            return self.buffer_docx
        except Exception as e:
            st.error(f"Error generando DOCX: {str(e)}")
            return None

    def generar_geojson(self):
        try:
            gdf_out = self.gdf.copy()
            res = self.resultados
            if res:
                gdf_out['area_ha'] = res.get('area_total_ha', 0)
                gdf_out['carbono_total_ton'] = res.get('carbono_total_ton', 0)
                gdf_out['shannon_promedio'] = res.get('shannon_promedio', 0)
                gdf_out['ecosistema'] = res.get('tipo_ecosistema', 'N/A')
                if 'analisis_forrajero' in res:
                    gdf_out['forraje_kg_ms_ha'] = res['analisis_forrajero']['disponibilidad_forrajera']['productividad_kg_ms_ha']
            geojson_str = gdf_out.to_json()
            return geojson_str
        except Exception as e:
            st.error(f"Error generando GeoJSON: {str(e)}")
            return json.dumps({"error": str(e)})

# ===============================
# FUNCIÓN PARA GENERAR INFORME CON IA
# ===============================
def generar_reporte_ia(resultados, gdf, sistema_mapas=None):
    import tempfile
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import datetime
    import io
    import os

    if not REPORTDOCX_AVAILABLE:
        st.error("python-docx no está instalado. No se puede generar el informe.")
        return None

    with tempfile.TemporaryDirectory() as tmpdir:
        doc = Document()
        section = doc.sections[0]
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        title = doc.add_heading('INFORME AMBIENTAL CON ANÁLISIS DE IA', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph(f'Fecha: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        df, stats = preparar_resumen(resultados)

        # 1. Resumen ejecutivo
        doc.add_heading('1. RESUMEN EJECUTIVO', level=1)
        tabla_resumen = doc.add_table(rows=1, cols=3)
        tabla_resumen.style = 'Light Shading'
        tabla_resumen.cell(0, 0).text = 'Métrica'
        tabla_resumen.cell(0, 1).text = 'Valor'
        tabla_resumen.cell(0, 2).text = 'Interpretación'

        metricas = [
            ('Área total', f"{stats['area_total_ha']:,.1f} ha", 'Superficie del área de estudio'),
            ('Carbono total', f"{stats['carbono_total_ton']:,.0f} ton C", 'Almacenamiento total de carbono'),
            ('CO₂ equivalente', f"{stats['co2_total_ton']:,.0f} ton CO₂e", 'Potencial de créditos de carbono'),
            ('Índice Shannon', f"{stats['shannon_promedio']:.3f}", 'Nivel de biodiversidad'),
            ('NDVI promedio', f"{stats['ndvi_promedio']:.3f}", 'Salud de la vegetación'),
            ('NDWI promedio', f"{stats['ndwi_promedio']:.3f}", 'Contenido de agua'),
            ('Tipo ecosistema', stats['tipo_ecosistema'], 'Vegetación predominante'),
            ('Puntos muestreo', str(stats['num_puntos']), 'Muestras analizadas')
        ]
        for i, (met, val, interp) in enumerate(metricas, 1):
            row = tabla_resumen.add_row().cells
            row[0].text = met
            row[1].text = val
            row[2].text = interp
        doc.add_paragraph()

        # 2. Análisis de Carbono
        doc.add_heading('2. ANÁLISIS DE CARBONO', level=1)
        if resultados.get('desglose_promedio'):
            doc.add_heading('Distribución por pools', level=2)
            tabla_pools = doc.add_table(rows=1, cols=3)
            tabla_pools.style = 'Light Shading'
            tabla_pools.cell(0, 0).text = 'Pool'
            tabla_pools.cell(0, 1).text = 'Descripción'
            tabla_pools.cell(0, 2).text = 'Ton C/ha'
            desc = {'AGB':'Biomasa Aérea Viva', 'BGB':'Biomasa de Raíces', 'DW':'Madera Muerta', 'LI':'Hojarasca', 'SOC':'Carbono Orgánico del Suelo'}
            for pool, valor in resultados['desglose_promedio'].items():
                row = tabla_pools.add_row().cells
                row[0].text = pool
                row[1].text = desc.get(pool, pool)
                row[2].text = f"{valor:.2f}"
            doc.add_paragraph()
            vis = Visualizaciones()
            fig_carbono = vis.crear_grafico_barras_carbono(resultados['desglose_promedio'])
            if fig_carbono:
                try:
                    img_bytes = fig_carbono.to_image(format='png', width=800, height=500, scale=2)
                    img_path = os.path.join(tmpdir, 'carbono.png')
                    with open(img_path, 'wb') as f:
                        f.write(img_bytes)
                    doc.add_picture(img_path, width=Inches(5))
                    doc.add_paragraph()
                except:
                    pass

        doc.add_heading('2.1 Interpretación técnica', level=2)
        analisis_carbono = generar_analisis_carbono(df, stats)
        doc.add_paragraph(analisis_carbono)

        # 3. Análisis de Biodiversidad
        doc.add_heading('3. ANÁLISIS DE BIODIVERSIDAD', level=1)
        if resultados.get('puntos_biodiversidad'):
            biodiv = resultados['puntos_biodiversidad'][0]
            tabla_biodiv = doc.add_table(rows=1, cols=2)
            tabla_biodiv.style = 'Light Shading'
            tabla_biodiv.cell(0, 0).text = 'Métrica'
            tabla_biodiv.cell(0, 1).text = 'Valor'
            metricas_bio = [
                ('Índice Shannon', f"{biodiv.get('indice_shannon', 0):.3f}"),
                ('Categoría', biodiv.get('categoria', 'N/A')),
                ('Riqueza de especies', str(biodiv.get('riqueza_especies', 0))),
                ('Abundancia total', f"{biodiv.get('abundancia_total', 0):,}")
            ]
            for met, val in metricas_bio:
                row = tabla_biodiv.add_row().cells
                row[0].text = met
                row[1].text = val
            doc.add_paragraph()
            fig_biodiv = vis.crear_grafico_radar_biodiversidad(biodiv)
            if fig_biodiv:
                try:
                    img_bytes = fig_biodiv.to_image(format='png', width=800, height=800, scale=2)
                    img_path = os.path.join(tmpdir, 'biodiv.png')
                    with open(img_path, 'wb') as f:
                        f.write(img_bytes)
                    doc.add_picture(img_path, width=Inches(5))
                    doc.add_paragraph()
                except:
                    pass

        doc.add_heading('3.1 Interpretación técnica', level=2)
        analisis_biodiv = generar_analisis_biodiversidad(df, stats)
        doc.add_paragraph(analisis_biodiv)

        # 4. Análisis de Índices Espectrales
        doc.add_heading('4. ANÁLISIS DE ÍNDICES ESPECTRALES', level=1)
        doc.add_heading('4.1 Interpretación técnica', level=2)
        analisis_espectral = generar_analisis_espectral(df, stats)
        doc.add_paragraph(analisis_espectral)

        # 5. Análisis Forrajero (NUEVO)
        doc.add_heading('5. ANÁLISIS FORRAJERO', level=1)
        if 'analisis_forrajero' in resultados:
            forrajero = resultados['analisis_forrajero']
            disp = forrajero['disponibilidad_forrajera']
            ev = forrajero['equivalentes_vaca']
            tabla_forraje = doc.add_table(rows=1, cols=2)
            tabla_forraje.style = 'Light Shading'
            tabla_forraje.cell(0, 0).text = 'Métrica'
            tabla_forraje.cell(0, 1).text = 'Valor'
            datos_f = [
                ('Productividad (kg MS/ha)', f"{disp['productividad_kg_ms_ha']:,.0f}"),
                ('Forraje aprovechable (ton)', f"{disp['forraje_aprovechable_kg_ms']/1000:.1f}"),
                ('EV por día', f"{ev['ev_por_dia']:.1f}"),
                ('EV recomendado (30 días)', f"{ev['ev_recomendado']:.1f}")
            ]
            for met, val in datos_f:
                row = tabla_forraje.add_row().cells
                row[0].text = met
                row[1].text = val
            doc.add_paragraph()
            fig_forrajero = vis.crear_grafico_forrajero(disp, ev)
            if fig_forrajero:
                try:
                    img_bytes = fig_forrajero.to_image(format='png', width=1000, height=700, scale=2)
                    img_path = os.path.join(tmpdir, 'forrajero.png')
                    with open(img_path, 'wb') as f:
                        f.write(img_bytes)
                    doc.add_picture(img_path, width=Inches(6))
                    doc.add_paragraph()
                except:
                    pass

        doc.add_heading('5.1 Interpretación técnica', level=2)
        analisis_forrajero = generar_analisis_forrajero(df, stats)  # función nueva
        doc.add_paragraph(analisis_forrajero)

        # 6. Mapas de calor
        if sistema_mapas:
            doc.add_heading('6. MAPAS DE CALOR CONTINUOS', level=1)
            variables = ['carbono', 'ndvi', 'ndwi', 'biodiversidad', 'forraje']
            titulos = ['Carbono (ton C/ha)', 'NDVI', 'NDWI', 'Biodiversidad (Shannon)', 'Productividad Forrajera (kg MS/ha)']
            for var, tit in zip(variables, titulos):
                mapa = sistema_mapas.crear_mapa_estatico(resultados, var, gdf)
                if mapa:
                    doc.add_heading(tit, level=2)
                    img_path = os.path.join(tmpdir, f'mapa_{var}.png')
                    with open(img_path, 'wb') as f:
                        f.write(mapa.getvalue())
                    doc.add_picture(img_path, width=Inches(6))
                    doc.add_paragraph()

        # 7. Recomendaciones Integradas
        doc.add_heading('7. RECOMENDACIONES DE MANEJO', level=1)
        recomendaciones = generar_recomendaciones_integradas(df, stats)
        doc.add_paragraph(recomendaciones)

        # 8. Metadatos
        doc.add_heading('8. METADATOS', level=1)
        metadatos = [
            ('Generado por', 'Sistema Satelital de Análisis Ambiental v3.0 con IA Gemini'),
            ('Fecha de generación', datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
            ('Número de puntos', str(stats['num_puntos']))
        ]
        for key, val in metadatos:
            p = doc.add_paragraph()
            p.add_run(f"{key}: ").bold = True
            p.add_run(val)

        docx_output = BytesIO()
        doc.save(docx_output)
        docx_output.seek(0)
        return docx_output

# ===============================
# FUNCIONES AUXILIARES
# ===============================
def validar_y_corregir_crs(gdf):
    if gdf is None or len(gdf) == 0:
        return gdf
    try:
        if gdf.crs is None:
            gdf = gdf.set_crs('EPSG:4326', inplace=False)
            st.info("ℹ️ Se asignó EPSG:4326 al archivo (no tenía CRS)")
        elif str(gdf.crs).upper() != 'EPSG:4326':
            original_crs = str(gdf.crs)
            gdf = gdf.to_crs('EPSG:4326')
            st.info(f"ℹ️ Transformado de {original_crs} a EPSG:4326")
        return gdf
    except Exception as e:
        st.warning(f"⚠️ Error al corregir CRS: {str(e)}")
        return gdf

def calcular_superficie(gdf):
    try:
        if gdf is None or len(gdf) == 0:
            return 0.0
        gdf = validar_y_corregir_crs(gdf)
        bounds = gdf.total_bounds
        if bounds[0] < -180 or bounds[2] > 180 or bounds[1] < -90 or bounds[3] > 90:
            st.warning("⚠️ Coordenadas fuera de rango para cálculo preciso de área")
            area_grados2 = gdf.geometry.area.sum()
            area_m2 = area_grados2 * 111000 * 111000
            return area_m2 / 10000
        gdf_projected = gdf.to_crs('EPSG:3857')
        area_m2 = gdf_projected.geometry.area.sum()
        return area_m2 / 10000
    except Exception as e:
        try:
            return gdf.geometry.area.sum() / 10000
        except:
            return 0.0

def dividir_poligono_en_cuadricula(poligono, puntos_forraje, n_celdas=100):
    try:
        bounds = poligono.bounds
        minx, miny, maxx, maxy = bounds
        n_cols = int(np.sqrt(n_celdas * (maxx - minx) / (maxy - miny)))
        n_rows = int(n_celdas / n_cols)
        if n_rows == 0:
            n_rows = 1
        width = (maxx - minx) / n_cols
        height = (maxy - miny) / n_rows
        celdas = []
        productividades = []
        for i in range(n_rows):
            for j in range(n_cols):
                cell_minx = minx + j * width
                cell_maxx = minx + (j + 1) * width
                cell_miny = miny + i * height
                cell_maxy = miny + (i + 1) * height
                cell_poly = Polygon([(cell_minx, cell_miny), (cell_maxx, cell_miny), (cell_maxx, cell_maxy), (cell_minx, cell_maxy)])
                intersection = poligono.intersection(cell_poly)
                if intersection.is_empty or intersection.area == 0:
                    continue
                puntos_dentro = [p['productividad_kg_ms_ha'] for p in puntos_forraje if Point(p['lon'], p['lat']).within(intersection)]
                if puntos_dentro:
                    prod_promedio = np.mean(puntos_dentro)
                else:
                    min_dist = float('inf')
                    prod_cercano = None
                    for p in puntos_forraje:
                        point = Point(p['lon'], p['lat'])
                        dist = intersection.distance(point)
                        if dist < min_dist:
                            min_dist = dist
                            prod_cercano = p['productividad_kg_ms_ha']
                    prod_promedio = prod_cercano if prod_cercano is not None else 0
                celdas.append(intersection)
                productividades.append(prod_promedio)
        gdf_celdas = gpd.GeoDataFrame({'geometry': celdas, 'productividad_kg_ms_ha': productividades}, crs='EPSG:4326')
        return gdf_celdas
    except Exception as e:
        st.warning(f"Error en dividir cuadrícula: {str(e)}")
        return gpd.GeoDataFrame()

def cargar_shapefile_desde_zip(zip_file):
    try:
        with tempfile.TemporaryDirectory() as tmp_dir:
            with zipfile.ZipFile(zip_file, 'r') as zip_ref:
                zip_ref.extractall(tmp_dir)
            shp_files = [f for f in os.listdir(tmp_dir) if f.endswith('.shp')]
            if shp_files:
                shp_path = os.path.join(tmp_dir, shp_files[0])
                gdf = gpd.read_file(shp_path)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
            else:
                st.error("❌ No se encontró ningún archivo .shp en el ZIP")
                return None
    except Exception as e:
        st.error(f"❌ Error cargando shapefile desde ZIP: {str(e)}")
        return None

def parsear_kml_manual(contenido_kml):
    try:
        root = ET.fromstring(contenido_kml)
        namespaces = {'kml': 'http://www.opengis.net/kml/2.2'}
        polygons = []
        for polygon_elem in root.findall('.//kml:Polygon', namespaces):
            coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
            if coords_elem is not None and coords_elem.text:
                coord_text = coords_elem.text.strip()
                coord_list = []
                for coord_pair in coord_text.split():
                    parts = coord_pair.split(',')
                    if len(parts) >= 2:
                        lon = float(parts[0])
                        lat = float(parts[1])
                        coord_list.append((lon, lat))
                if len(coord_list) >= 3:
                    polygons.append(Polygon(coord_list))
        if not polygons:
            for multi_geom in root.findall('.//kml:MultiGeometry', namespaces):
                for polygon_elem in multi_geom.findall('.//kml:Polygon', namespaces):
                    coords_elem = polygon_elem.find('.//kml:coordinates', namespaces)
                    if coords_elem is not None and coords_elem.text:
                        coord_text = coords_elem.text.strip()
                        coord_list = []
                        for coord_pair in coord_text.split():
                            parts = coord_pair.split(',')
                            if len(parts) >= 2:
                                lon = float(parts[0])
                                lat = float(parts[1])
                                coord_list.append((lon, lat))
                        if len(coord_list) >= 3:
                            polygons.append(Polygon(coord_list))
        if polygons:
            gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
            return gdf
        else:
            for placemark in root.findall('.//kml:Placemark', namespaces):
                for elem_name in ['Polygon', 'LineString', 'Point', 'LinearRing']:
                    elem = placemark.find(f'.//kml:{elem_name}', namespaces)
                    if elem is not None:
                        coords_elem = elem.find('.//kml:coordinates', namespaces)
                        if coords_elem is not None and coords_elem.text:
                            coord_text = coords_elem.text.strip()
                            coord_list = []
                            for coord_pair in coord_text.split():
                                parts = coord_pair.split(',')
                                if len(parts) >= 2:
                                    lon = float(parts[0])
                                    lat = float(parts[1])
                                    coord_list.append((lon, lat))
                            if len(coord_list) >= 3:
                                polygons.append(Polygon(coord_list))
                            break
        if polygons:
            gdf = gpd.GeoDataFrame({'geometry': polygons}, crs='EPSG:4326')
            return gdf
        return None
    except Exception as e:
        st.error(f"❌ Error parseando KML manualmente: {str(e)}")
        return None

def cargar_kml(kml_file):
    try:
        if kml_file.name.endswith('.kmz'):
            with tempfile.TemporaryDirectory() as tmp_dir:
                with zipfile.ZipFile(kml_file, 'r') as zip_ref:
                    zip_ref.extractall(tmp_dir)
                kml_files = [f for f in os.listdir(tmp_dir) if f.endswith('.kml')]
                if kml_files:
                    kml_path = os.path.join(tmp_dir, kml_files[0])
                    with open(kml_path, 'r', encoding='utf-8') as f:
                        contenido = f.read()
                    gdf = parsear_kml_manual(contenido)
                    if gdf is not None:
                        return gdf
                    else:
                        try:
                            gdf = gpd.read_file(kml_path)
                            gdf = validar_y_corregir_crs(gdf)
                            return gdf
                        except:
                            st.error("❌ No se pudo cargar el archivo KML/KMZ")
                            return None
                else:
                    st.error("❌ No se encontró ningún archivo .kml en el KMZ")
                    return None
        else:
            contenido = kml_file.read().decode('utf-8')
            gdf = parsear_kml_manual(contenido)
            if gdf is not None:
                return gdf
            else:
                kml_file.seek(0)
                gdf = gpd.read_file(kml_file)
                gdf = validar_y_corregir_crs(gdf)
                return gdf
    except Exception as e:
        st.error(f"❌ Error cargando archivo KML/KMZ: {str(e)}")
        return None

def cargar_archivo_parcela(uploaded_file):
    try:
        if uploaded_file.name.endswith('.zip'):
            gdf = cargar_shapefile_desde_zip(uploaded_file)
        elif uploaded_file.name.endswith(('.kml', '.kmz')):
            gdf = cargar_kml(uploaded_file)
        elif uploaded_file.name.endswith('.geojson'):
            gdf = gpd.read_file(uploaded_file)
            gdf = validar_y_corregir_crs(gdf)
        else:
            st.error("❌ Formato de archivo no soportado")
            return None

        if gdf is not None:
            gdf = validar_y_corregir_crs(gdf)
            gdf = gdf.explode(ignore_index=True)
            gdf = gdf[gdf.geometry.geom_type.isin(['Polygon', 'MultiPolygon'])]
            if len(gdf) == 0:
                st.error("❌ No se encontraron polígonos en el archivo")
                return None
            geometria_unida = gdf.unary_union
            gdf_unido = gpd.GeoDataFrame([{'geometry': geometria_unida}], crs='EPSG:4326')
            gdf_unido = validar_y_corregir_crs(gdf_unido)
            st.info(f"✅ Se unieron {len(gdf)} polígono(s) en una sola geometría.")
            gdf_unido['id_zona'] = 1
            return gdf_unido
        return gdf
    except Exception as e:
        st.error(f"❌ Error cargando archivo: {str(e)}")
        import traceback
        st.error(f"Detalle: {traceback.format_exc()}")
        return None

# ===============================
# FUNCIÓN PRINCIPAL DE ANÁLISIS
# ===============================
def ejecutar_analisis_completo(gdf, tipo_ecosistema, num_puntos, usar_gee=False):
    try:
        area_total = calcular_superficie(gdf)
        poligono = gdf.geometry.iloc[0]
        bounds = poligono.bounds

        clima = ConectorClimaticoTropical()
        verra = MetodologiaVerra()
        biodiversidad = AnalisisBiodiversidad()
        forrajero = AnalisisForrajero()

        if tipo_ecosistema in ['pampa', 'seco']:
            sistema_forrajero = 'pastizal_natural'
        elif tipo_ecosistema in ['amazonia', 'choco']:
            sistema_forrajero = 'silvopastoril'
        else:
            sistema_forrajero = 'pastizal_natural'

        puntos_carbono = []
        puntos_biodiversidad = []
        puntos_ndvi = []
        puntos_ndwi = []
        puntos_ndre = []
        puntos_msavi = []
        puntos_evi = []
        puntos_forraje = []

        carbono_total = 0
        co2_total = 0
        shannon_promedio = 0
        ndvi_promedio = 0
        ndwi_promedio = 0
        area_por_punto = max(area_total / num_puntos, 0.1)

        puntos_generados = 0
        max_intentos = num_puntos * 10

        while puntos_generados < num_puntos and len(puntos_carbono) < max_intentos:
            lat = bounds[1] + random.random() * (bounds[3] - bounds[1])
            lon = bounds[0] + random.random() * (bounds[2] - bounds[0])
            point = Point(lon, lat)

            if poligono.contains(point):
                datos_clima = clima.obtener_datos_climaticos(lat, lon)
                ndvi = 0.5 + random.uniform(-0.2, 0.3)
                base_ndwi = 0.1
                if datos_clima['precipitacion'] > 2000:
                    base_ndwi += 0.3
                elif datos_clima['precipitacion'] < 800:
                    base_ndwi -= 0.2
                ndwi = base_ndwi + random.uniform(-0.2, 0.2)
                ndwi = max(-0.5, min(0.8, ndwi))
                ndre = min(1.0, max(-1.0, ndvi * 0.95 + random.uniform(-0.05, 0.1)))
                msavi = min(1.0, max(0.0, ndvi * 0.85 + random.uniform(-0.1, 0.05)))
                evi = min(1.0, max(0.0, ndvi * 1.2 + random.uniform(-0.1, 0.1)))

                carbono_info = verra.calcular_carbono_hectarea(ndvi, tipo_ecosistema, datos_clima['precipitacion'])
                biodiv_info = biodiversidad.calcular_shannon(ndvi, tipo_ecosistema, area_por_punto, datos_clima['precipitacion'])
                forraje_info = forrajero.estimar_disponibilidad_forrajera(ndvi, sistema_forrajero, area_por_punto)

                carbono_total += carbono_info['carbono_total_ton_ha'] * area_por_punto
                co2_total += carbono_info['co2_equivalente_ton_ha'] * area_por_punto
                shannon_promedio += biodiv_info['indice_shannon']
                ndvi_promedio += ndvi
                ndwi_promedio += ndwi

                puntos_carbono.append({'lat': lat, 'lon': lon, 'carbono_ton_ha': carbono_info['carbono_total_ton_ha'], 'ndvi': ndvi, 'precipitacion': datos_clima['precipitacion']})
                biodiv_info['lat'] = lat
                biodiv_info['lon'] = lon
                puntos_biodiversidad.append(biodiv_info)
                puntos_ndvi.append({'lat': lat, 'lon': lon, 'ndvi': ndvi})
                puntos_ndwi.append({'lat': lat, 'lon': lon, 'ndwi': ndwi})
                puntos_ndre.append({'lat': lat, 'lon': lon, 'ndre': ndre})
                puntos_msavi.append({'lat': lat, 'lon': lon, 'msavi': msavi})
                puntos_evi.append({'lat': lat, 'lon': lon, 'evi': evi})
                puntos_forraje.append({'lat': lat, 'lon': lon, 'productividad_kg_ms_ha': forraje_info['productividad_kg_ms_ha']})

                puntos_generados += 1

        if puntos_generados > 0:
            shannon_promedio /= puntos_generados
            ndvi_promedio /= puntos_generados
            ndwi_promedio /= puntos_generados

        carbono_promedio = verra.calcular_carbono_hectarea(ndvi_promedio, tipo_ecosistema, 1500)

        # Análisis forrajero
        disponibilidad_forrajera = forrajero.estimar_disponibilidad_forrajera(ndvi_promedio, sistema_forrajero, area_total)
        equivalentes_vaca = forrajero.calcular_equivalentes_vaca(disponibilidad_forrajera['forraje_aprovechable_kg_ms'], dias_permanencia=30)
        sublotes = forrajero.dividir_lote_en_sublotes(area_total, disponibilidad_forrajera['productividad_kg_ms_ha'], heterogeneidad=0.3)
        gdf_cuadricula = dividir_poligono_en_cuadricula(poligono, puntos_forraje, n_celdas=200)

        resultados = {
            'area_total_ha': area_total,
            'carbono_total_ton': round(carbono_total, 2),
            'co2_total_ton': round(co2_total, 2),
            'carbono_promedio_ha': round(carbono_total / area_total, 2) if area_total > 0 else 0,
            'shannon_promedio': round(shannon_promedio, 3),
            'ndvi_promedio': round(ndvi_promedio, 3),
            'ndwi_promedio': round(ndwi_promedio, 3),
            'puntos_carbono': puntos_carbono,
            'puntos_biodiversidad': puntos_biodiversidad,
            'puntos_ndvi': puntos_ndvi,
            'puntos_ndwi': puntos_ndwi,
            'puntos_ndre': puntos_ndre,
            'puntos_msavi': puntos_msavi,
            'puntos_evi': puntos_evi,
            'puntos_forraje': puntos_forraje,
            'gdf_cuadricula': gdf_cuadricula,
            'tipo_ecosistema': tipo_ecosistema,
            'num_puntos': puntos_generados,
            'desglose_promedio': carbono_promedio['desglose'] if carbono_promedio else {},
            'usar_gee': usar_gee,
            'analisis_forrajero': {
                'sistema_forrajero': sistema_forrajero,
                'disponibilidad_forrajera': disponibilidad_forrajera,
                'equivalentes_vaca': equivalentes_vaca,
                'sublotes': sublotes,
                'forrajero': forrajero
            }
        }
        return resultados
    except Exception as e:
        st.error(f"Error en ejecutar_analisis_completo: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
        return None

# ===============================
# FUNCIONES DE VISUALIZACIÓN
# ===============================
def mostrar_mapas_calor():
    st.header("🗺️ Mapas de Calor Continuos")
    if st.session_state.poligono_data is None:
        st.info("Ejecute el análisis primero.")
        return

    tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
        "🌍 Área Base", "🌳 Carbono", "📈 NDVI", "💧 NDWI", "🦋 Biodiversidad", "🌿 Forrajero", "🎭 Combinado"
    ])

    with tab1:
        st.subheader("Mapa Base del Área de Estudio")
        if st.session_state.mapa:
            folium_static(st.session_state.mapa, width=1000, height=650)
        else:
            st.info("No hay mapa base.")

    sistema = SistemaMapas()
    with tab2:
        if 'puntos_carbono' in st.session_state.resultados:
            mapa = sistema.crear_mapa_calor_interpolado(st.session_state.resultados, 'carbono', st.session_state.poligono_data)
            if mapa:
                folium_static(mapa, width=1000, height=650)
            else:
                st.warning("No se pudo generar el mapa.")
    with tab3:
        if 'puntos_ndvi' in st.session_state.resultados:
            mapa = sistema.crear_mapa_calor_interpolado(st.session_state.resultados, 'ndvi', st.session_state.poligono_data)
            if mapa:
                folium_static(mapa, width=1000, height=650)
    with tab4:
        if 'puntos_ndwi' in st.session_state.resultados:
            mapa = sistema.crear_mapa_calor_interpolado(st.session_state.resultados, 'ndwi', st.session_state.poligono_data)
            if mapa:
                folium_static(mapa, width=1000, height=650)
    with tab5:
        if 'puntos_biodiversidad' in st.session_state.resultados:
            mapa = sistema.crear_mapa_calor_interpolado(st.session_state.resultados, 'biodiversidad', st.session_state.poligono_data)
            if mapa:
                folium_static(mapa, width=1000, height=650)
    with tab6:
        if 'puntos_forraje' in st.session_state.resultados:
            mapa = sistema.crear_mapa_calor_interpolado(st.session_state.resultados, 'forraje', st.session_state.poligono_data)
            if mapa:
                folium_static(mapa, width=1000, height=650)
    with tab7:
        st.info("Mapa combinado con control de capas (implementación pendiente)")

def mostrar_dashboard():
    st.header("📊 Dashboard Ejecutivo")
    if st.session_state.resultados is None:
        st.info("Ejecute el análisis primero.")
        return
    res = st.session_state.resultados
    html_kpi = Visualizaciones.crear_metricas_kpi(
        res.get('carbono_total_ton', 0),
        res.get('co2_total_ton', 0),
        res.get('shannon_promedio', 0),
        res.get('area_total_ha', 0)
    )
    st.markdown(html_kpi, unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("📈 NDVI promedio", f"{res.get('ndvi_promedio', 0):.3f}")
    with col2:
        st.metric("💧 NDWI promedio", f"{res.get('ndwi_promedio', 0):.3f}")
    with col3:
        st.metric("🎯 Puntos analizados", res.get('num_puntos', 0))

    if 'analisis_forrajero' in res:
        st.subheader("🐮 Métricas Forrajeras")
        col_f1, col_f2, col_f3 = st.columns(3)
        with col_f1:
            st.metric("Productividad", f"{res['analisis_forrajero']['disponibilidad_forrajera']['productividad_kg_ms_ha']:,.0f}", "kg MS/ha")
        with col_f2:
            st.metric("EV Recomendados", f"{res['analisis_forrajero']['equivalentes_vaca']['ev_recomendado']:.1f}")
        with col_f3:
            st.metric("Sublotes", len(res['analisis_forrajero']['sublotes']))

    col1, col2 = st.columns(2)
    with col1:
        fig_carbono = Visualizaciones.crear_grafico_barras_carbono(res.get('desglose_promedio', {}))
        if fig_carbono:
            st.plotly_chart(fig_carbono, use_container_width=True)
    with col2:
        if res.get('puntos_biodiversidad'):
            fig_biodiv = Visualizaciones.crear_grafico_radar_biodiversidad(res['puntos_biodiversidad'][0])
            if fig_biodiv:
                st.plotly_chart(fig_biodiv, use_container_width=True)

def mostrar_carbono():
    st.header("🌳 Análisis de Carbono")
    if st.session_state.resultados is None:
        st.info("Ejecute el análisis primero.")
        return
    res = st.session_state.resultados
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Carbono Total", f"{res.get('carbono_total_ton', 0):,.0f} ton C")
    with col2:
        st.metric("Potencial Créditos", f"{res.get('co2_total_ton', 0)/1000:,.1f} k")
    with col3:
        valor_economico = res.get('co2_total_ton', 0) * 15
        st.metric("Valor Aprox.", f"${valor_economico:,.0f} USD")
    if res.get('desglose_promedio'):
        st.subheader("Distribución por Pools")
        df_pools = pd.DataFrame({
            'Pool': list(res['desglose_promedio'].keys()),
            'Ton C/ha': list(res['desglose_promedio'].values())
        })
        st.dataframe(df_pools, use_container_width=True)

def mostrar_biodiversidad():
    st.header("🦋 Análisis de Biodiversidad")
    if st.session_state.resultados is None:
        st.info("Ejecute el análisis primero.")
        return
    res = st.session_state.resultados
    if res.get('puntos_biodiversidad'):
        biodiv = res['puntos_biodiversidad'][0]
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Índice Shannon", f"{biodiv.get('indice_shannon', 0):.3f}", biodiv.get('categoria', ''))
        with col2:
            st.metric("Riqueza especies", f"{biodiv.get('riqueza_especies', 0)}")
        with col3:
            st.metric("Abundancia total", f"{biodiv.get('abundancia_total', 0):,}")
        # Gráfico de distribución
        shannon_vals = [p.get('indice_shannon', 0) for p in res['puntos_biodiversidad']]
        fig = go.Figure(data=[go.Histogram(x=shannon_vals, nbinsx=15, marker_color='#8b5cf6')])
        fig.update_layout(title='Distribución del Índice de Shannon', xaxis_title='Valor', yaxis_title='Frecuencia', height=400)
        st.plotly_chart(fig, use_container_width=True)

def mostrar_analisis_forrajero():
    st.header("🐮 Análisis Forrajero")
    if st.session_state.resultados is None or 'analisis_forrajero' not in st.session_state.resultados:
        st.info("Ejecute el análisis completo primero.")
        return
    res = st.session_state.resultados
    forrajero_data = res['analisis_forrajero']
    disp = forrajero_data['disponibilidad_forrajera']
    ev = forrajero_data['equivalentes_vaca']

    st.subheader("🌿 Disponibilidad Forrajera")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Productividad", f"{disp['productividad_kg_ms_ha']:,.0f}", "kg MS/ha")
    with col2:
        st.metric("Disponible total", f"{disp['disponibilidad_total_kg_ms']/1000:,.1f}", "ton MS")
    with col3:
        st.metric("Aprovechable", f"{disp['forraje_aprovechable_kg_ms']/1000:,.1f}", "ton MS")
    with col4:
        st.metric("Categoría", disp['categoria_productividad'].title())

    st.subheader("🐄 Equivalentes Vaca")
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("EV por día", f"{ev['ev_por_dia']:.1f}")
    with col2:
        st.metric("EV para 30 días", f"{ev['ev_para_periodo']:.1f}")
    with col3:
        st.metric("EV recomendado", f"{ev['ev_recomendado']:.1f}")

    if 'sublotes' in forrajero_data and forrajero_data['sublotes']:
        st.subheader("📋 Sublotes")
        df_sub = pd.DataFrame(forrajero_data['sublotes'])
        st.dataframe(df_sub, use_container_width=True, hide_index=True)

    # Gráfico forrajero
    fig_forrajero = Visualizaciones.crear_grafico_forrajero(disp, ev)
    if fig_forrajero:
        st.plotly_chart(fig_forrajero, use_container_width=True)

    # Mapa de sublotes (coroplético)
    if 'gdf_cuadricula' in res and not res['gdf_cuadricula'].empty:
        st.subheader("🗺️ Mapa de Productividad por Sublotes")
        sistema = SistemaMapas()
        try:
            bounds = st.session_state.poligono_data.total_bounds
            centro = [(bounds[1] + bounds[3]) / 2, (bounds[0] + bounds[2]) / 2]
            m = folium.Map(location=centro, zoom_start=12, tiles='OpenStreetMap')
            min_prod = res['gdf_cuadricula']['productividad_kg_ms_ha'].min()
            max_prod = res['gdf_cuadricula']['productividad_kg_ms_ha'].max()
            colormap = LinearColormap(colors=['#8B4513', '#CD853F', '#F4A460', '#9ACD32', '#32CD32', '#006400'], vmin=min_prod, vmax=max_prod)
            colormap.caption = 'Productividad Forrajera (kg MS/ha)'
            folium.GeoJson(
                res['gdf_cuadricula'],
                style_function=lambda feature: {
                    'fillColor': colormap(feature['properties']['productividad_kg_ms_ha']),
                    'color': 'black',
                    'weight': 0.5,
                    'fillOpacity': 0.7
                },
                tooltip=folium.GeoJsonTooltip(fields=['productividad_kg_ms_ha'], aliases=['Productividad:'], localize=True)
            ).add_to(m)
            folium.GeoJson(
                st.session_state.poligono_data.geometry.iloc[0],
                style_function=lambda x: {'fillColor': 'transparent', 'color': '#1d4ed8', 'weight': 3, 'dashArray': '5, 5'}
            ).add_to(m)
            colormap.add_to(m)
            folium_static(m, width=1000, height=600)
        except Exception as e:
            st.warning(f"No se pudo generar el mapa de sublotes: {str(e)}")

    # Calculadora interactiva
    with st.expander("📊 Calculadora de Equivalentes Vaca"):
        num_ev_input = st.number_input("Número de EV disponibles:", min_value=1.0, max_value=1000.0, value=50.0, step=1.0)
        dias_input = st.number_input("Días de permanencia deseada:", min_value=1, max_value=365, value=30, step=1)
        if st.button("Calcular días"):
            forrajero = forrajero_data['forrajero']
            dias_calc = forrajero.calcular_dias_permanencia(disp['forraje_aprovechable_kg_ms'], num_ev_input)
            st.success(f"**Resultado:** {num_ev_input:.0f} EV pueden pastar {dias_calc['dias_recomendados']} días")
            col1, col2, col3 = st.columns(3)
            with col1: st.metric("Días básicos", f"{dias_calc['dias_basico']:.1f}")
            with col2: st.metric("Días ajustados", f"{dias_calc['dias_ajustado']:.1f}")
            with col3: st.metric("Recomendados", dias_calc['dias_recomendados'])

def mostrar_comparacion():
    st.header("📈 Análisis Comparativo")
    if st.session_state.resultados is None:
        st.info("Ejecute el análisis primero.")
        return
    res = st.session_state.resultados
    if all(k in res for k in ['puntos_carbono', 'puntos_ndvi', 'puntos_ndwi', 'puntos_biodiversidad']):
        fig = Visualizaciones.crear_grafico_comparativo(
            res['puntos_carbono'], res['puntos_ndvi'], res['puntos_ndwi'], res['puntos_biodiversidad']
        )
        if fig:
            st.plotly_chart(fig, use_container_width=True)
    # Correlaciones
    st.subheader("🔗 Correlaciones")
    try:
        n = min(100, len(res['puntos_carbono']))
        carbono_vals = [p['carbono_ton_ha'] for p in res['puntos_carbono'][:n]]
        ndvi_vals = [p['ndvi'] for p in res['puntos_ndvi'][:n]]
        ndwi_vals = [p['ndwi'] for p in res['puntos_ndwi'][:n]]
        shannon_vals = [p['indice_shannon'] for p in res['puntos_biodiversidad'][:n]]
        corr1 = np.corrcoef(carbono_vals, ndvi_vals)[0,1] if len(carbono_vals)>1 else 0
        corr2 = np.corrcoef(carbono_vals, shannon_vals)[0,1] if len(carbono_vals)>1 else 0
        corr3 = np.corrcoef(ndvi_vals, shannon_vals)[0,1] if len(ndvi_vals)>1 else 0
        corr4 = np.corrcoef(ndwi_vals, shannon_vals)[0,1] if len(ndwi_vals)>1 else 0
        col1, col2, col3, col4 = st.columns(4)
        with col1: st.metric("C vs NDVI", f"{corr1:.3f}")
        with col2: st.metric("C vs Shannon", f"{corr2:.3f}")
        with col3: st.metric("NDVI vs Shannon", f"{corr3:.3f}")
        with col4: st.metric("NDWI vs Shannon", f"{corr4:.3f}")
    except Exception as e:
        st.warning(f"No se pudieron calcular correlaciones: {str(e)}")

def mostrar_informe():
    st.header("📥 Informe Completo")
    if st.session_state.resultados is None or st.session_state.poligono_data is None:
        st.info("Ejecute el análisis primero.")
        return

    st.markdown("### Generar informe con todos los análisis")
    sistema = SistemaMapas()
    generador = GeneradorReportes(st.session_state.resultados, st.session_state.poligono_data, sistema)

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if REPORTPDF_AVAILABLE:
            if st.button("📄 Generar PDF", use_container_width=True):
                pdf = generador.generar_pdf()
                if pdf:
                    st.download_button("⬇️ Descargar PDF", pdf, f"informe_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf", "application/pdf")
    with col2:
        if REPORTDOCX_AVAILABLE:
            if st.button("📘 Generar DOCX", use_container_width=True):
                docx = generador.generar_docx()
                if docx:
                    st.download_button("⬇️ Descargar DOCX", docx, f"informe_{datetime.now().strftime('%Y%m%d_%H%M')}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col3:
        if st.button("🤖 Generar Informe con IA", use_container_width=True):
            with st.spinner("Generando informe con IA..."):
                reporte_ia = generar_reporte_ia(st.session_state.resultados, st.session_state.poligono_data, sistema)
                if reporte_ia:
                    st.download_button("⬇️ Descargar Informe IA", reporte_ia, f"informe_IA_{datetime.now().strftime('%Y%m%d_%H%M')}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with col4:
        if st.button("🌍 Generar GeoJSON", use_container_width=True):
            geojson = generador.generar_geojson()
            if geojson:
                st.download_button("⬇️ Descargar GeoJSON", geojson, f"area_{datetime.now().strftime('%Y%m%d_%H%M')}.geojson", "application/geo+json")

# ===============================
# MAIN
# ===============================
def main():
    if 'gee_authenticated' not in st.session_state:
        st.session_state.gee_authenticated = False
        st.session_state.gee_project = ''
        if GEE_AVAILABLE:
            inicializar_gee()
    if 'poligono_data' not in st.session_state:
        st.session_state.poligono_data = None
    if 'resultados' not in st.session_state:
        st.session_state.resultados = None
    if 'mapa' not in st.session_state:
        st.session_state.mapa = None

    st.title("🌎 Sistema Satelital de Análisis Ambiental Integral")
    st.markdown("### Carbono + Biodiversidad + Análisis Forrajero")

    with st.sidebar:
        st.header("📁 Carga de Datos")
        if GEE_AVAILABLE and st.session_state.gee_authenticated:
            st.success(f"✅ GEE Conectado")
        uploaded_file = st.file_uploader("Cargar polígono (KML, GeoJSON, SHP, KMZ)", type=['kml', 'geojson', 'zip', 'kmz'])
        if uploaded_file:
            with st.spinner("Procesando archivo..."):
                gdf = cargar_archivo_parcela(uploaded_file)
                if gdf is not None:
                    st.session_state.poligono_data = gdf
                    area_ha = calcular_superficie(gdf)
                    st.info(f"📍 Área calculada: {area_ha:,.1f} ha")
                    sistema = SistemaMapas()
                    st.session_state.mapa = sistema.crear_mapa_area(gdf, zoom_auto=True)

        if st.session_state.poligono_data is not None:
            st.header("⚙️ Configuración")
            tipo_ecosistema = st.selectbox("Tipo de ecosistema", ['amazonia', 'choco', 'andes', 'pampa', 'seco', 'cultivo', 'vid', 'agricola'])
            num_puntos = st.slider("Número de puntos de muestreo", 10, 200, 50)
            usar_gee = False
            if GEE_AVAILABLE and st.session_state.gee_authenticated:
                usar_gee = st.checkbox("Usar datos reales de GEE")
            if st.button("🚀 Ejecutar Análisis Completo", type="primary", use_container_width=True):
                with st.spinner("Analizando..."):
                    resultados = ejecutar_analisis_completo(st.session_state.poligono_data, tipo_ecosistema, num_puntos, usar_gee)
                    if resultados:
                        st.session_state.resultados = resultados
                        st.success("✅ Análisis completado!")

    if st.session_state.poligono_data is None:
        st.info("👈 Cargue un polígono en el panel lateral para comenzar")
        with st.expander("📋 Información del Sistema"):
            st.markdown("""
            ### Sistema Integrado de Análisis Ambiental Satelital
            **Características:**
            - 🌳 Metodología Verra VCS para cálculo de carbono
            - 🦋 Índice de Shannon para biodiversidad
            - 📈 NDVI, NDWI, NDRE, MSAVI, EVI
            - 🐮 Análisis forrajero (productividad, EV, rotación)
            - 🗺️ Mapas de calor continuos con interpolación KNN
            - 📊 Dashboard interactivo y gráficos
            - 📄 Informes PDF/DOCX/GeoJSON y con IA (Gemini)
            """)
    else:
        tabs = st.tabs(["🗺️ Mapas", "📊 Dashboard", "🌳 Carbono", "🦋 Biodiversidad", "🐮 Forrajero", "📈 Comparación", "📥 Informe"])
        with tabs[0]: mostrar_mapas_calor()
        with tabs[1]: mostrar_dashboard()
        with tabs[2]: mostrar_carbono()
        with tabs[3]: mostrar_biodiversidad()
        with tabs[4]: mostrar_analisis_forrajero()
        with tabs[5]: mostrar_comparacion()
        with tabs[6]: mostrar_informe()

if __name__ == "__main__":
    main()
